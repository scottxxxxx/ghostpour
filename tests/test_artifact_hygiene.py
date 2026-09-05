"""Dashes and arrows inside generated files, and in the chat normalizer."""

import zipfile
from io import BytesIO

from app.services.artifact_hygiene import DOCX, PPTX, XLSX, scrub_office_text
from app.services.text_hygiene import normalize_dashes


def _docx_with_dashes() -> bytes:
    import docx
    d = docx.Document()
    d.add_heading("Austin Bike Mechanics — Automated Repair Service", 0)
    t = d.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "Data availability — production"
    t.rows[0].cells[1].text = "October–December 2026"
    d.add_paragraph("Sept 10 preview → Nov 2026 GA")
    d.sections[0].footer.paragraphs[0].text = "Confidential — for internal distribution"
    buf = BytesIO(); d.save(buf); return buf.getvalue()


def _docx_text(blob: bytes) -> str:
    import docx
    d = docx.Document(BytesIO(blob))
    parts = [p.text for p in d.paragraphs]
    parts += [c.text for t in d.tables for r in t.rows for c in r.cells]
    parts += [p.text for s in d.sections for p in s.footer.paragraphs]
    return "\n".join(parts)


def test_the_live_docx_shape_comes_out_clean_and_readable():
    """The 2026-09-05 file: title joined with an em dash, a status table
    with em dashes, a month range with an en dash, a footer, an arrow."""
    blob = _docx_with_dashes()
    assert "—" in _docx_text(blob)
    out, counts = scrub_office_text(blob, DOCX)
    text = _docx_text(out)
    assert "—" not in text and "–" not in text and "→" not in text
    assert "Austin Bike Mechanics, Automated Repair Service" in text
    assert "Data availability, production" in text
    assert "October to December 2026" in text          # a tight en dash between words is a range
    assert "Sept 10 preview to Nov 2026 GA" in text    # an arrow reads as "to"
    assert "Confidential, for internal distribution" in text
    assert counts["dashes"] == 4 and counts["arrows"] == 1 and counts["parts"] == 2


def test_an_xlsx_shared_string_and_inline_string_are_rewritten():
    import openpyxl
    wb = openpyxl.Workbook(); ws = wb.active
    ws["A1"] = "CTS deployment — CAB approval"; ws["A2"] = "Q3–Q4"; ws["A3"] = 12
    buf = BytesIO(); wb.save(buf)
    out, counts = scrub_office_text(buf.getvalue(), XLSX)
    wb2 = openpyxl.load_workbook(BytesIO(out))
    assert wb2.active["A1"].value == "CTS deployment, CAB approval"
    assert wb2.active["A2"].value == "Q3 to Q4"  # a tight en dash between words is a range
    assert wb2.active["A3"].value == 12
    assert counts["dashes"] == 2


def test_a_pptx_slide_text_run_is_rewritten():
    buf = BytesIO()
    with zipfile.ZipFile(buf, "w") as z:
        z.writestr("[Content_Types].xml", "<Types/>")
        z.writestr("ppt/slides/slide1.xml", '<p:sld><a:t>Roadmap — Q4</a:t><a:t>plain</a:t></p:sld>')
    out, counts = scrub_office_text(buf.getvalue(), PPTX)
    assert b"Roadmap, Q4" in zipfile.ZipFile(BytesIO(out)).read("ppt/slides/slide1.xml")
    assert counts["dashes"] == 1


def test_clean_files_and_unknown_mimes_pass_through_byte_for_byte():
    import docx
    d = docx.Document(); d.add_paragraph("No dashes here, just commas."); buf = BytesIO(); d.save(buf)
    blob = buf.getvalue()
    assert scrub_office_text(blob, DOCX) == (blob, {"dashes": 0, "arrows": 0, "parts": 0})
    assert scrub_office_text(b"%PDF-1.4 \xe2\x80\x94 not an office file", "application/pdf")[0].startswith(b"%PDF")


def test_a_corrupt_zip_returns_the_original_bytes():
    blob = b"not a zip at all \xe2\x80\x94"
    assert scrub_office_text(blob, DOCX)[0] == blob


def test_arrows_in_chat_text_read_as_to():
    assert normalize_dashes("Timeline (Sept 10 preview → Nov 2026 GA)") == "Timeline (Sept 10 preview to Nov 2026 GA)"
    assert normalize_dashes("A ⇒ B") == "A to B"
    assert normalize_dashes("no glyphs") == "no glyphs"


def test_the_build_leg_tells_the_model_the_rule_applies_inside_files():
    src = open("app/services/providers/anthropic.py").read()
    i = src.index("Inside every file you create")
    assert "em dash" in src[i:i + 400] and "arrow" in src[i:i + 400]


def test_the_collector_scrubs_after_the_word_rebuild_and_logs_the_generation_id():
    src = open("app/services/document_generation.py").read()
    assert src.index("rebuild_docx, content)") < src.index("scrub_office_text, content, mime)") < src.index("staging.stage(")
    assert "artifact_dashes_rewritten generation_id=%s" in src
    chat = open("app/routers/chat.py").read()
    i = chat.index("generated_payload = await collect_generated_files(")
    assert "generation_id=_generation_id" in chat[i:i + 500]
