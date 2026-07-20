"""Document generation phase 2a: staging store, serve endpoint, gate,
adapter arming, artifact collection, response wire field.

Design: docs/design/documents-phase2-returned-files.md. Ships dark
(documents.generation.enabled false); allowed_users (shared with phase 1)
is the e2e lane.
"""

import json
import re

import pytest

from app.services.document_generation import (
    _walk_file_ids,
    generation_gate,
    load_generation_config,
)

XLSX = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


def _configs(gen=None, allowed=None):
    docs = {"enabled": False, "allowed_users": allowed or []}
    if gen is not None:
        docs["generation"] = gen
    return {"client-config": {"documents": docs}}


# --- config + gate ---

def test_generation_defaults_ship_dark():
    cfg = load_generation_config({})
    assert cfg["enabled"] is False
    assert cfg["min_tier"] == "pro"
    assert len(cfg["formats"]) == 4  # all four launch formats (§9 decision 2)
    assert cfg["max_files_out"] == 2 and cfg["max_file_out_mb"] == 25


def test_bundled_config_generation_live_pro_gated():
    # FLIPPED LIVE 2026-07-13 (Scott: TestFlight-scale user base, wants the
    # feature in hand this week). Load-bearing invariants: pro gate intact,
    # locales agree, all four formats.
    for f in ("client-config.json", "client-config.es.json", "client-config.ja.json"):
        gen = json.load(open(f"config/remote/{f}"))["documents"]["generation"]
        assert gen["enabled"] is True and gen["min_tier"] == "pro"
        assert len(gen["formats"]) == 4


def _gate(**over):
    kw = dict(
        remote_configs=_configs(gen={"enabled": True, "min_tier": "pro"}),
        tier_name="pro", managed_routing=True, provider="anthropic",
        prompt_mode="ProjectChat", user_identity={"u1"},
    )
    kw.update(over)
    return generation_gate(**kw)


def test_gate_matrix():
    assert _gate() is True
    assert _gate(prompt_mode="PostMeetingChat") is True
    # every mechanical requirement individually closes the gate
    assert _gate(prompt_mode=None) is False
    assert _gate(prompt_mode="InterviewScorecard") is False
    assert _gate(managed_routing=False) is False
    assert _gate(provider="openrouter") is False
    # enabled+tier path
    assert _gate(tier_name="plus") is False
    assert _gate(remote_configs=_configs(gen={"enabled": False})) is False
    # allowed_users overrides enabled AND tier (e2e lane, shared w/ phase 1)
    assert _gate(
        remote_configs=_configs(gen={"enabled": False}, allowed=["scott@x.com"]),
        tier_name="plus", user_identity={"scott@x.com"},
    ) is True


# --- adapter arming ---

def _adapter():
    from app.services.providers.anthropic import AnthropicAdapter
    return AnthropicAdapter(
        api_key="test", base_url="https://example.invalid/v1/messages",
        auth_header="x-api-key", auth_prefix="",
    )


def _body(generation):
    from app.models.chat import ChatRequest
    return ChatRequest(
        provider="anthropic", model="claude-sonnet-4-6",
        system_prompt="sys", user_content="make me a spreadsheet",
        generation=generation, max_tokens=4096,
    )


def test_adapter_arms_generation():
    api_body, headers = _adapter()._build_body(_body(True))
    skills = {s["skill_id"] for s in api_body["container"]["skills"]}
    assert skills == {"xlsx", "pptx", "docx", "pdf"}
    assert any(t["type"] == "code_execution_20260521" for t in api_body["tools"])
    assert api_body["max_tokens"] >= 16000
    # spike-mandated: cache breakpoint on the user content ($1.04 -> $0.33)
    text_parts = [p for p in api_body["messages"][0]["content"] if p["type"] == "text"]
    assert text_parts[-1]["cache_control"] == {"type": "ephemeral"}
    assert "code-execution-2025-08-25" in headers["anthropic-beta"]
    assert "skills-2025-10-02" in headers["anthropic-beta"]


def test_adapter_unarmed_is_untouched():
    api_body, headers = _adapter()._build_body(_body(False))
    assert "container" not in api_body
    assert "tools" not in api_body
    assert api_body["max_tokens"] == 4096
    assert "anthropic-beta" not in headers or "skills" not in headers.get("anthropic-beta", "")


# --- artifact collection ---

def test_walk_file_ids_finds_and_dedups():
    raw = json.dumps({"content": [
        {"type": "text", "text": "done"},
        {"type": "bash_code_execution_tool_result", "content": {
            "type": "bash_code_execution_result",
            "content": [{"type": "bash_code_execution_output", "file_id": "file_A"}]}},
        {"type": "bash_code_execution_tool_result", "content": {
            "type": "bash_code_execution_result",
            "content": [{"type": "bash_code_execution_output", "file_id": "file_A"},
                        {"type": "bash_code_execution_output", "file_id": "file_B"}]}},
    ]})
    assert _walk_file_ids(raw) == ["file_A", "file_B"]
    assert _walk_file_ids("not json") == []
    assert _walk_file_ids(json.dumps({"content": [{"type": "text", "text": "hi"}]})) == []


@pytest.mark.asyncio
async def test_collect_downloads_validates_and_stages(tmp_path, monkeypatch):
    import aiosqlite
    from app.services import generated_files as staging
    from app.services.document_generation import collect_generated_files

    monkeypatch.setattr(staging, "STAGING_DIR", tmp_path / "gen")

    class _Resp:
        def __init__(self, status, payload=None, content=b""):
            self.status_code, self._p, self.content = status, payload, content
        def json(self):
            return self._p

    calls = {}
    class _Client:
        def __init__(self, *a, **k): pass
        async def __aenter__(self): return self
        async def __aexit__(self, *a): return False
        async def get(self, url, headers=None):
            calls.setdefault("urls", []).append(url)
            if url.endswith("/content"):
                return _Resp(200, content=b"PK\x03\x04 fake xlsx bytes")
            if url.endswith("file_big"):
                return _Resp(200, {"filename": "huge.xlsx", "mime_type": XLSX,
                                   "size_bytes": 99 * 1024 * 1024})
            if url.endswith("file_txt"):
                return _Resp(200, {"filename": "notes.txt", "mime_type": "text/plain",
                                   "size_bytes": 10})
            return _Resp(200, {"filename": "tracker.xlsx", "mime_type": XLSX,
                               "size_bytes": 20})
    import app.services.document_generation as dg
    monkeypatch.setattr(dg.httpx, "AsyncClient", _Client)

    raw = json.dumps({"content": [
        {"type": "bash_code_execution_tool_result", "content": {
            "content": [{"type": "bash_code_execution_output", "file_id": "file_ok"},
                        {"type": "bash_code_execution_output", "file_id": "file_big"},
                        {"type": "bash_code_execution_output", "file_id": "file_txt"}]}},
    ]})

    async with aiosqlite.connect(":memory:") as db:
        db.row_factory = aiosqlite.Row
        await db.execute("""CREATE TABLE generated_files (
            id TEXT PRIMARY KEY, user_id TEXT NOT NULL, app_id TEXT,
            name TEXT NOT NULL, media_type TEXT NOT NULL, size_bytes INTEGER NOT NULL,
            storage_path TEXT NOT NULL, created_at TEXT NOT NULL, expires_at TEXT NOT NULL)""")
        out = await collect_generated_files(
            db, raw_response_json=raw, api_key="k",
            remote_configs=_configs(gen={"enabled": True}),
            user_id="u1", app_id="shouldersurf",
        )
    # max_files_out=2 truncates to [file_ok, file_big]; big skipped on size;
    # only the valid xlsx stages
    assert len(out) == 1
    g = out[0]
    assert g["name"] == "tracker.xlsx" and g["media_type"] == XLSX
    assert g["url"].startswith("/v1/generated-files/gpf_")
    assert g["size_bytes"] == len(b"PK\x03\x04 fake xlsx bytes")
    import hashlib
    assert g["sha256"] == hashlib.sha256(b"PK\x03\x04 fake xlsx bytes").hexdigest()


# --- staging store semantics ---

@pytest.mark.asyncio
async def test_staging_expiry_ownership_and_cap(tmp_path, monkeypatch):
    import aiosqlite
    from app.services import generated_files as staging

    monkeypatch.setattr(staging, "STAGING_DIR", tmp_path / "gen")
    async with aiosqlite.connect(":memory:") as db:
        db.row_factory = aiosqlite.Row
        await db.execute("""CREATE TABLE generated_files (
            id TEXT PRIMARY KEY, user_id TEXT NOT NULL, app_id TEXT,
            name TEXT NOT NULL, media_type TEXT NOT NULL, size_bytes INTEGER NOT NULL,
            storage_path TEXT NOT NULL, created_at TEXT NOT NULL, expires_at TEXT NOT NULL)""")

        row = await staging.stage(db, user_id="u1", app_id="ss",
                                  name="a.xlsx", media_type=XLSX, content=b"x" * 100)
        assert row and row["expires_at"] > row["url"]  # sanity: fields present
        # owner fetch works; other user 404-equivalent
        assert await staging.fetch(db, row["file_id"], "u1") is not None
        assert await staging.fetch(db, row["file_id"], "someone-else") is None
        # live cap: a stage that would exceed 50MB is refused, not an error
        monkeypatch.setattr(staging, "PER_USER_LIVE_CAP_BYTES", 150)
        assert await staging.stage(db, user_id="u1", app_id="ss",
                                   name="b.xlsx", media_type=XLSX, content=b"y" * 100) is None
        # expiry: force-expire then purge deletes row + bytes
        await db.execute("UPDATE generated_files SET expires_at = '2000-01-01T00:00:00+00:00'")
        await db.commit()
        assert await staging.fetch(db, row["file_id"], "u1") is None
        n = await staging.purge_expired(db)
        assert n == 1
        left = await (await db.execute("SELECT COUNT(*) AS n FROM generated_files")).fetchone()
        assert left["n"] == 0


# --- serve endpoint (via app test client) ---

def test_serve_endpoint_auth_ownership_expiry(client, pro_user, tmp_db_path, tmp_path, monkeypatch):
    import sqlite3
    from datetime import datetime, timedelta, timezone

    blob = tmp_path / "gpf_test"
    blob.write_bytes(b"PK\x03\x04 artifact")
    future = (datetime.now(timezone.utc) + timedelta(hours=1)).isoformat()
    past = (datetime.now(timezone.utc) - timedelta(hours=1)).isoformat()
    con = sqlite3.connect(tmp_db_path)
    uid = pro_user["user_id"] if "user_id" in pro_user else None
    if uid is None:
        row = con.execute("SELECT id FROM users WHERE email LIKE 'test-pro-user%'").fetchone()
        uid = row[0]
    con.execute("INSERT INTO generated_files VALUES (?,?,?,?,?,?,?,?,?)",
                ("gpf_live", uid, "shouldersurf", "t.xlsx", XLSX, 11, str(blob), future, future))
    con.execute("INSERT INTO generated_files VALUES (?,?,?,?,?,?,?,?,?)",
                ("gpf_dead", uid, "shouldersurf", "t.xlsx", XLSX, 11, str(blob), past, past))
    con.commit(); con.close()

    h = pro_user["headers"]
    r = client.get("/v1/generated-files/gpf_live", headers=h)
    assert r.status_code == 200
    assert r.content == b"PK\x03\x04 artifact"
    assert "no-store" in r.headers.get("cache-control", "")
    assert client.get("/v1/generated-files/gpf_dead", headers=h).status_code == 404
    assert client.get("/v1/generated-files/gpf_missing", headers=h).status_code == 404
    assert client.get("/v1/generated-files/gpf_live").status_code in (401, 403, 422)  # no auth


# --- generation reliability: per-leg timeout + no OR fallback ---

@pytest.mark.asyncio
@pytest.mark.parametrize("generation,expected_timeout", [(True, 400.0), (False, None)])
async def test_generation_leg_gets_extended_timeout(monkeypatch, generation, expected_timeout):
    adapter = _adapter()
    seen = {}

    async def fake_post(url, body, headers, timeout=None):
        seen["timeout"] = timeout
        return 200, {"content": [{"type": "text", "text": "ok"}],
                     "stop_reason": "end_turn", "usage": {}}, "", ""
    monkeypatch.setattr(adapter, "_post", fake_post)

    await adapter.send_request(_body(generation))
    assert seen["timeout"] == expected_timeout


@pytest.mark.asyncio
async def test_generation_turn_never_falls_back_to_or(monkeypatch):
    import httpx as _httpx
    from unittest.mock import AsyncMock, MagicMock
    from app.services.anthropic_or_fallback import route_with_fallback

    router = MagicMock()
    router.route = AsyncMock(side_effect=_httpx.ReadTimeout("leg timed out"))
    settings = MagicMock()

    with pytest.raises(_httpx.ReadTimeout):
        await route_with_fallback(router, _body(True), db=None, settings=settings)
    # one attempt, no OR retry — a fallback would mean a second route() call
    assert router.route.await_count == 1


@pytest.mark.asyncio
async def test_non_generation_turn_still_falls_back(monkeypatch):
    import httpx as _httpx
    from unittest.mock import AsyncMock, MagicMock
    import app.services.anthropic_or_fallback as fb

    ok = MagicMock(name="or_response")
    router = MagicMock()
    router.route = AsyncMock(side_effect=[_httpx.ReadTimeout("boom"), ok])
    monkeypatch.setattr(fb, "_alert_on_fallback", AsyncMock())

    out = await fb.route_with_fallback(router, _body(False), db=None, settings=MagicMock())
    assert out is ok
    assert router.route.await_count == 2


# --- chat bubble = closing summary, not working transcript ---

@pytest.mark.asyncio
@pytest.mark.parametrize("generation,expected_text", [
    (True, "Done — your tracker is attached."),                 # last block only
    (False, "Let me think.\nDone — your tracker is attached."),  # unchanged join
])
async def test_generation_bubble_is_closing_summary(monkeypatch, generation, expected_text):
    adapter = _adapter()

    async def fake_post(url, body, headers, timeout=None):
        return 200, {"content": [
            {"type": "text", "text": "Let me think."},
            {"type": "server_tool_use", "id": "t1", "name": "bash_code_execution", "input": {}},
            {"type": "bash_code_execution_tool_result", "content": {}},
            {"type": "text", "text": "Done — your tracker is attached."},
        ], "stop_reason": "end_turn", "usage": {}}, "", ""
    monkeypatch.setattr(adapter, "_post", fake_post)

    resp = await adapter.send_request(_body(generation))
    assert resp.text == expected_text


def test_raw_provider_payloads_never_reach_the_wire(client, free_user, mock_provider):
    from tests.conftest import chat_request
    r = client.post("/v1/chat", json=chat_request(), headers=free_user["headers"])
    body = r.json()
    assert "raw_request_json" not in body
    assert "raw_response_json" not in body
    assert "text" in body and "model" in body  # normal shape intact


# --- confirmation envelope (handoff Part 1) ---

def test_confirmation_defaults_dark_but_bundles_live_and_agree():
    # Code DEFAULTS stay dark (a missing/stale config must never enable the
    # flow); the BUNDLES ship enabled since 2026-07-11 — while
    # generation.enabled is false this reaches only the allowed_users lane
    # (the gate runs before the confirmation logic).
    from app.services.document_generation import load_generation_config
    conf = load_generation_config({})["confirmation"]
    assert conf["enabled"] is False
    assert conf["expected_seconds"] == 150 and conf["poll_after_seconds"] == 5
    assert set(conf["format_nouns"]) == {"xlsx", "docx", "pptx", "pdf"}
    # re-lit 2026-07-12 on SS's client-ready signal (offer rendering live);
    # still reaches only the allowed_users lane while generation.enabled is
    # false. Policy: ALWAYS ASK (Scott).
    for f in ("client-config.json", "client-config.es.json", "client-config.ja.json"):
        c = json.load(open(f"config/remote/{f}"))["documents"]["generation"]["confirmation"]
        assert c["enabled"] is True
        assert "{format}" in c["offer_text"]
        assert set(c["format_nouns"]) == {"xlsx", "docx", "pptx", "pdf"}


def test_locale_variant_supplies_envelope_text():
    from app.services.document_generation import load_generation_config
    configs = {
        "client-config": {"documents": {"generation": {"confirmation": {"offer_text": "EN {format}"}}}},
        "client-config.es": {"documents": {"generation": {"confirmation": {"offer_text": "ES {format}"}}}},
    }
    assert load_generation_config(configs)["confirmation"]["offer_text"] == "EN {format}"
    assert load_generation_config(configs, locale="es")["confirmation"]["offer_text"] == "ES {format}"
    # unknown locale falls back to base
    assert load_generation_config(configs, locale="fr")["confirmation"]["offer_text"] == "EN {format}"


def test_build_offer_envelope_shape():
    from app.services.document_generation import _CONFIRMATION_DEFAULTS, build_offer_envelope
    env = build_offer_envelope(_CONFIRMATION_DEFAULTS, "docx")
    fs = env["feature_state"]
    assert fs["feature"] == "document_generation"
    assert fs["state"] == "confirmation_required"
    cta = fs["cta"]
    assert cta["kind"] == "generation_offer" and cta["action"] == "confirm_generation"
    assert "a native Word document (.docx)" in cta["text"]
    assert "two minutes" in cta["text"]                 # sets the wait expectation
    assert "right here in chat" in cta["text"]          # offers the inline alternative
    assert cta["details"] == {"expected_format": "docx", "expected_seconds": 150, "gist": ""}
    # unknown/None format degrades to the default noun path, never a KeyError
    assert build_offer_envelope(_CONFIRMATION_DEFAULTS, None)["feature_state"]["cta"]["details"]["expected_format"] == "xlsx"


@pytest.mark.asyncio
async def test_classifier_fail_open_and_strict_parse():
    from unittest.mock import AsyncMock, MagicMock
    from app.services.document_generation import classify_generation_intent

    # provider error -> None (fail-open)
    router = MagicMock()
    router.route = AsyncMock(side_effect=RuntimeError("boom"))
    assert await classify_generation_intent(router, "make me a spreadsheet") is None

    # junk text -> None
    router.route = AsyncMock(return_value=MagicMock(text="I think probably yes?"))
    assert await classify_generation_intent(router, "make me a spreadsheet") is None

    # valid YES with prose wrapping -> parsed, format normalized
    router.route = AsyncMock(return_value=MagicMock(
        text='Sure: {"file_request": true, "format": "docx"} done'))
    out = await classify_generation_intent(router, "write this up as a word doc")
    assert out == {"file_request": True, "format": "docx", "gist": ""}

    # valid NO -> file_request False
    router.route = AsyncMock(return_value=MagicMock(
        text='{"file_request": false, "format": null}'))
    # (input mentions "report" so the prefilter admits it; the classifier
    # still says no — prefilter is recall-biased, classifier decides)
    out = await classify_generation_intent(router, "what did we decide about the report?")
    assert out == {"file_request": False, "format": None, "gist": ""}

    # bogus format value normalizes to None rather than leaking to the wire
    router.route = AsyncMock(return_value=MagicMock(
        text='{"file_request": true, "format": "exe"}'))
    out = await classify_generation_intent(router, "make me a file")
    assert out == {"file_request": True, "format": None, "gist": ""}


@pytest.mark.asyncio
async def test_classifier_meters_via_on_subcall():
    from unittest.mock import AsyncMock, MagicMock
    from app.services.document_generation import classify_generation_intent

    router = MagicMock()
    router.route = AsyncMock(return_value=MagicMock(
        text='{"file_request": true, "format": "xlsx"}'))
    seen = {}
    async def subcall(creq, cresp, cms):
        seen["call_type"] = creq.get_meta("call_type")
        seen["model"] = creq.model
    out = await classify_generation_intent(router, "build a tracker", on_subcall=subcall)
    assert out["file_request"] is True
    assert seen["call_type"] == "generation_intent"
    assert seen["model"].startswith("claude-haiku")


# --- transport Phase A (e2e through the app) ---

def _enable_confirmed_generation(client):
    docs = client.app.state.remote_configs["client-config"].setdefault("documents", {})
    docs["generation"] = {"enabled": True, "min_tier": "free",
                          "confirmation": {"enabled": True, "expected_seconds": 150}}


def test_confirmed_turn_rides_sse_and_records_rescue_row(client, free_user, mock_provider, tmp_db_path):
    import json as _json
    import sqlite3
    from tests.conftest import chat_request

    _enable_confirmed_generation(client)
    r = client.post("/v1/chat", json=chat_request(
        prompt_mode="ProjectChat", call_type="query",
        metadata={"generation_confirmed": True, "generation_id": "gen-sse-e2e"},
        user_content="Build me a tracking spreadsheet",
    ), headers=free_user["headers"])
    assert r.status_code == 200
    assert r.headers["content-type"].startswith("text/event-stream")
    assert "event: generation_started" in r.text
    assert "event: generation_result" in r.text
    started = _json.loads(
        r.text.split("event: generation_started\ndata: ")[1].split("\n")[0])
    assert started["expected_seconds"] == 150
    result = _json.loads(
        r.text.split("event: generation_result\ndata: ")[1].split("\n")[0])
    assert result["text"]                       # the normal JSON body, verbatim
    assert "raw_request_json" not in result     # wire-slim holds on SSE too

    con = sqlite3.connect(tmp_db_path)
    row = con.execute("SELECT status, user_id FROM generations "
                      "WHERE generation_id='gen-sse-e2e'").fetchone()
    con.close()
    assert row is not None and row[0] == "done"


def test_unconfirmed_soft_turn_fails_open_to_normal_json(client, free_user, mock_provider):
    from tests.conftest import chat_request
    _enable_confirmed_generation(client)
    # confirmation live + no confirmed flag + SOFT phrasing (explicit verbs
    # now short-circuit to a guaranteed offer): the classifier runs against
    # the mock provider, whose reply parses as junk -> fail-open -> normal
    # chat with the soft-intent TEASER attached, never a dead turn.
    r = client.post("/v1/chat", json=chat_request(
        prompt_mode="ProjectChat", call_type="query",
        user_content="what did we say about the spreadsheet yesterday?",
    ), headers=free_user["headers"])
    assert r.status_code == 200
    assert r.headers["content-type"].startswith("application/json")
    body = r.json()
    assert body["text"]
    assert body["feature_state"]["cta"]["kind"] == "generation_teaser"
    assert body.get("generated_files") is None or "generated_files" not in body


# --- docx Word-compat rebuild (backstop) ---

def _make_source_docx():
    import io
    import docx
    d = docx.Document()
    d.add_heading("Onboarding Guide", level=1)
    p = d.add_paragraph()
    r = p.add_run("Welcome ")
    r2 = p.add_run("aboard")
    r2.bold = True
    d.add_paragraph("First item", style="List Bullet")
    d.add_paragraph("Second item", style="List Bullet")
    t = d.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "Owner"
    t.rows[0].cells[1].text = "Task"
    t.rows[1].cells[0].text = "Scott"
    t.rows[1].cells[1].text = "Review"
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def test_docx_rebuild_preserves_content_and_produces_valid_docx():
    import io
    import docx
    from app.services.docx_rebuild import rebuild_docx
    rebuilt = rebuild_docx(_make_source_docx())
    d = docx.Document(io.BytesIO(rebuilt))
    texts = [p.text for p in d.paragraphs]
    assert "Onboarding Guide" in texts
    assert any("Welcome aboard" == t for t in texts)
    heading = next(p for p in d.paragraphs if p.text == "Onboarding Guide")
    assert heading.style.name == "Title"  # leading heading promoted to Title
    bullets = [p for p in d.paragraphs if p.text in ("First item", "Second item")]
    assert all(p.style.name == "List Bullet" for p in bullets)
    bold_run = next(r for p in d.paragraphs for r in p.runs if r.text == "aboard")
    assert bold_run.bold is True
    assert len(d.tables) == 1
    assert d.tables[0].rows[1].cells[0].text == "Scott"


def test_docx_rebuild_fails_open_on_garbage():
    from app.services.docx_rebuild import rebuild_docx
    garbage = b"not a docx at all"
    assert rebuild_docx(garbage) == garbage


def test_adapter_steers_docx_toolchain_on_generation():
    api_body, _ = _adapter()._build_body(_body(True))
    sys_texts = " ".join(b.get("text", "") for b in api_body["system"])
    assert "python-docx" in sys_texts and "docx.js" in sys_texts
    # non-generation turns get no steering
    api_body2, _ = _adapter()._build_body(_body(False))
    sys_texts2 = " ".join(b.get("text", "") for b in api_body2["system"])
    assert "python-docx" not in sys_texts2


def test_docx_rebuild_checklist_glyphs_drop_the_bullet():
    import io
    import docx
    from app.services.docx_rebuild import rebuild_docx
    d = docx.Document()
    d.add_paragraph("\u2610 Request ServiceNow access", style="List Bullet")
    d.add_paragraph("Regular bullet item", style="List Bullet")
    buf = io.BytesIO(); d.save(buf)
    out = docx.Document(io.BytesIO(rebuild_docx(buf.getvalue())))
    check = next(p for p in out.paragraphs if "ServiceNow" in p.text)
    plain = next(p for p in out.paragraphs if "Regular bullet" in p.text)
    assert check.style.name == "Normal"        # glyph is the marker
    assert plain.style.name == "List Bullet"   # real bullets keep the style


def test_docx_rebuild_title_header_and_page_footer():
    import io
    import docx
    from app.services.docx_rebuild import rebuild_docx
    d = docx.Document()
    d.add_heading("Quarterly Plan", level=1)
    d.add_heading("Goals", level=2)
    d.add_paragraph("Ship phase two.")
    buf = io.BytesIO(); d.save(buf)
    out = docx.Document(io.BytesIO(rebuild_docx(buf.getvalue())))
    assert out.paragraphs[0].style.name == "Title"
    assert out.paragraphs[0].text == "Quarterly Plan"
    goals = next(p for p in out.paragraphs if p.text == "Goals")
    assert goals.style.name == "Heading 2"          # section headings untouched
    sec = out.sections[0]
    assert sec.header.paragraphs[0].text == "Quarterly Plan"
    footer_xml = sec.footer.paragraphs[0]._p.xml
    assert 'w:instr=" PAGE "' in footer_xml and 'w:instr=" NUMPAGES "' in footer_xml


# --- conversational confirmation (Part 1 v2) ---

def test_offer_store_one_shot_and_ttl(monkeypatch):
    from app.services import generation_offers as go
    oid = go.create("u1", "docx", "for onboarding")
    assert go.take("u2", oid) is None            # not yours
    offer = go.take("u1", oid)
    assert offer == {"format": "docx", "gist": "for onboarding",
                     "template_id": None, "ask_content": "",
                     "images": [], "images_dropped": False,
                     "expires": offer["expires"]}
    assert go.take("u1", oid) is None            # one-shot: dead after a reply
    oid2 = go.create("u1", "xlsx", "")
    monkeypatch.setattr(go, "OFFER_TTL_S", 0)
    key = ("u1", oid2)
    go._OFFERS[key]["expires"] -= 9999
    assert go.take("u1", oid2) is None           # expired


def test_offer_store_keeps_originating_images_with_cap():
    """2026-07-19 fabricated-spreadsheet incident: reply sends carry chat
    history only, so an image-sourced build must inherit the ORIGINATING
    photo from the offer — and over-cap images are dropped WITH a marker
    so the arming path can refuse to generate blind."""
    from app.services import generation_offers as go
    oid = go.create("u1", "xlsx", "from image", images=["aW1nMQ==", "aW1nMg=="])
    offer = go.take("u1", oid)
    assert offer["images"] == ["aW1nMQ==", "aW1nMg=="]
    assert offer["images_dropped"] is False

    big = "x" * (go._IMAGES_CAP_CHARS + 1)
    oid2 = go.create("u1", "xlsx", "huge", images=[big])
    offer2 = go.take("u1", oid2)
    assert offer2["images"] == []
    assert offer2["images_dropped"] is True


def test_ask_references_images_marker():
    from app.services.document_generation import ask_references_images
    assert ask_references_images(
        "Project: X\n[1 image(s) attached for visual context]\nGive me an "
        "excel file from this image")
    assert ask_references_images("[3 image(s) attached for visual context]")
    assert not ask_references_images("make me a spreadsheet of Q3 sales")
    assert not ask_references_images("")


def test_generation_monthly_cap_parse():
    from app.services.document_generation import generation_monthly_cap
    cfgs = {"tiers": {"tiers": {
        "pro": {"feature_definitions": {"generation": {"generations_per_month": 100}}},
        "plus": {"feature_definitions": {"generation": {"generations_per_month": None}}},
        "free": {"feature_definitions": {}},
    }}}
    assert generation_monthly_cap(cfgs, "pro") == 100
    assert generation_monthly_cap(cfgs, "plus") is None      # explicit null = uncapped
    assert generation_monthly_cap(cfgs, "free") is None      # no block = uncapped
    assert generation_monthly_cap(cfgs, "unknown") is None
    assert generation_monthly_cap({}, "pro") is None


def test_generation_cap_quietly_disarms_everything(client, free_user, mock_provider,
                                                   tmp_db_path):
    """Quiet monthly count cap (2026-07-19): at cap, a confirmed explicit
    file ask neither arms generation nor draws an offer — the turn is a
    plain chat answer. No CTA, no error, no counter surface."""
    import sqlite3
    from tests.conftest import chat_request

    _enable_confirmed_generation(client)
    client.app.state.remote_configs.setdefault("tiers", {"tiers": {}})[
        "tiers"].setdefault("free", {}).setdefault(
        "feature_definitions", {})["generation"] = {"generations_per_month": 0}

    r = client.post("/v1/chat", json=chat_request(
        prompt_mode="ProjectChat", call_type="query",
        metadata={"generation_confirmed": True, "generation_id": "gen-capped"},
        user_content="Build me a tracking spreadsheet",
    ), headers=free_user["headers"])
    assert r.status_code == 200
    assert r.headers["content-type"].startswith("application/json")   # not SSE
    assert "feature_state" not in r.json()                            # no offer/teaser
    con = sqlite3.connect(tmp_db_path)
    # the id-bearing turn still records a terminal row (rescue-resolution
    # contract) — but as a plain chat answer: no files staged, no count
    row = con.execute(
        "SELECT files_json FROM generations WHERE generation_id='gen-capped'").fetchone()
    used = con.execute(
        "SELECT generations_used FROM users WHERE email LIKE 'test-free-user%'").fetchone()[0]
    con.close()
    assert row is not None and row[0] == "[]"
    assert used == 0
    # honest at-cap notice rode the routed request's system prompt
    routed = mock_provider.call_args[0][0]
    assert "FILE GENERATION NOTICE" in (routed.system_prompt or "")
    assert "resets on" in routed.system_prompt


def test_generations_used_counts_only_file_producing_builds(client, free_user,
                                                            mock_provider, tmp_db_path,
                                                            monkeypatch):
    """generations_used counts builds that staged artifacts (driven here
    through the template lane, which stages a real file). Plain chat
    turns carrying a rescue generation_id record done rows too (rescue
    resolution) — those must NOT count. The increment lives in the
    router, not generation_turns.finish, so the terminal-record service
    stays usable against the bare generations DDL."""
    import json as _json
    import sqlite3
    from unittest.mock import AsyncMock
    import app.services.document_generation as dg
    from app.services import generation_offers as go
    from tests.conftest import chat_request

    _enable_confirmed_generation(client)
    uid = free_user.get("user_id")
    if uid is None:
        con = sqlite3.connect(tmp_db_path)
        uid = con.execute("SELECT id FROM users WHERE email LIKE 'test-free-user%'").fetchone()[0]
        con.close()

    plan = {"project": "Count Test", "tasks": [
        {"id": 1, "name": "Phase", "type": "phase", "parent_id": None,
         "owner": None, "status": "in_progress",
         "start": "2026-07-01", "end": "2026-07-05", "depends_on": []},
        {"id": 2, "name": "Do the thing", "type": "task", "parent_id": 1,
         "owner": "Scott", "status": "in_progress",
         "start": "2026-07-01", "end": "2026-07-05", "depends_on": []},
    ]}
    mock_provider.return_value = mock_provider.canned_response.model_copy(
        update={"text": _json.dumps(plan)})

    oid = go.create(uid, "xlsx", "gantt", template_id="gantt_smartsheet",
                    ask_content="build a gantt chart for the plan")
    monkeypatch.setattr(dg, "interpret_offer_reply",
                        AsyncMock(return_value={"confirm": True, "format": "xlsx"}))

    r = client.post("/v1/chat", json=chat_request(
        prompt_mode="ProjectChat", call_type="query",
        metadata={"offer_id": oid, "generation_id": "gen-count-tmpl"},
        user_content="Yes",
    ), headers=free_user["headers"])
    assert r.status_code == 200
    assert "event: generation_result" in r.text

    con = sqlite3.connect(tmp_db_path)
    used = con.execute("SELECT generations_used FROM users WHERE id = ?", (uid,)).fetchone()[0]
    files = con.execute("SELECT files_json FROM generations WHERE generation_id='gen-count-tmpl'").fetchone()[0]
    con.close()
    assert files != "[]"     # the build staged a real artifact
    assert used == 1


def test_dashboard_edits_generation_cap_and_it_enforces(client):
    """The generic tier-field tunable writes generation.generations_per_month
    (lockstep locales, hot reload) — the dashboard's Save button IS the
    enforcement change."""
    from app.services.document_generation import generation_monthly_cap
    r = client.put("/webhooks/admin/tunable/tier-field",
                   json={"tier": "pro", "feature": "generation",
                         "field": "generations_per_month", "value": 7},
                   headers={"X-Admin-Key": "test-admin-key"})
    assert r.status_code == 200
    assert generation_monthly_cap(client.app.state.remote_configs, "pro") == 7
    # clearing = uncapped
    r2 = client.put("/webhooks/admin/tunable/tier-field",
                    json={"tier": "pro", "feature": "generation",
                          "field": "generations_per_month", "value": None},
                    headers={"X-Admin-Key": "test-admin-key"})
    assert r2.status_code == 200
    assert generation_monthly_cap(client.app.state.remote_configs, "pro") is None


def test_offer_envelope_conversational_with_gist_and_offer_id():
    from app.services.document_generation import _CONFIRMATION_DEFAULTS, build_offer_envelope
    env = build_offer_envelope(_CONFIRMATION_DEFAULTS, "docx",
                               gist="for onboarding new people", offer_id="abc123")
    cta = env["feature_state"]["cta"]
    assert cta["text"] == ("Sounds like you want a native Word document "
                           "(.docx) for onboarding new people. Building the "
                           "real file takes about two minutes, or I can just "
                           "lay it out right here in chat. Want the file?")
    assert cta["details"]["offer_id"] == "abc123"
    assert cta["details"]["gist"] == "for onboarding new people"
    # no gist -> falls back to the original template, no dangling braces
    env2 = build_offer_envelope(_CONFIRMATION_DEFAULTS, "xlsx", gist="", offer_id="x")
    assert "{" not in env2["feature_state"]["cta"]["text"]


@pytest.mark.asyncio
async def test_interpreter_confirms_declines_and_revises():
    from unittest.mock import AsyncMock, MagicMock
    from app.services.document_generation import interpret_offer_reply
    offer = {"format": "docx", "gist": "for onboarding"}

    router = MagicMock()
    router.route = AsyncMock(return_value=MagicMock(
        text='{"confirm": true, "format": null}'))
    out = await interpret_offer_reply(router, offer, "yes go ahead")
    assert out == {"confirm": True, "format": "docx"}      # offered format kept

    router.route = AsyncMock(return_value=MagicMock(
        text='{"confirm": true, "format": "xlsx"}'))
    out = await interpret_offer_reply(router, offer, "actually make it a spreadsheet")
    assert out == {"confirm": True, "format": "xlsx"}      # revised intent

    router.route = AsyncMock(return_value=MagicMock(
        text='{"confirm": false, "format": null}'))
    out = await interpret_offer_reply(router, offer, "no, what time is the standup?")
    assert out["confirm"] is False

    router.route = AsyncMock(side_effect=RuntimeError("boom"))
    out = await interpret_offer_reply(router, offer, "yes")
    assert out["confirm"] is False                          # fail-open = not armed


def test_chat_confirm_arms_generation_on_the_reply_turn(client, free_user, mock_provider,
                                                        tmp_db_path, monkeypatch):
    import sqlite3
    from unittest.mock import AsyncMock
    import app.services.document_generation as dg
    from app.services import generation_offers as go
    from tests.conftest import chat_request

    _enable_confirmed_generation(client)
    uid = free_user.get("user_id")
    if uid is None:
        con = sqlite3.connect(tmp_db_path)
        uid = con.execute("SELECT id FROM users WHERE email LIKE 'test-free-user%'").fetchone()[0]
        con.close()
    oid = go.create(uid, "docx", "for onboarding")
    monkeypatch.setattr(dg, "interpret_offer_reply",
                        AsyncMock(return_value={"confirm": True, "format": "docx"}))

    r = client.post("/v1/chat", json=chat_request(
        prompt_mode="ProjectChat", call_type="query",
        metadata={"offer_id": oid, "generation_id": "gen-chat-yes"},
        user_content="yes, go ahead",
    ), headers=free_user["headers"])
    assert r.status_code == 200
    assert r.headers["content-type"].startswith("text/event-stream")
    assert "event: generation_result" in r.text
    con = sqlite3.connect(tmp_db_path)
    row = con.execute("SELECT status FROM generations WHERE generation_id='gen-chat-yes'").fetchone()
    con.close()
    assert row and row[0] == "done"
    assert go.take(uid, oid) is None            # offer consumed


def test_typed_yes_inherits_offer_images_and_arms(client, free_user, mock_provider,
                                                  tmp_db_path, monkeypatch):
    """Image-sourced ask + stored images: the confirmed turn inherits the
    photo from the offer and generation proceeds (2026-07-19 incident,
    fixed half: the build sees the real image)."""
    import sqlite3
    from unittest.mock import AsyncMock
    import app.services.document_generation as dg
    from app.services import generation_offers as go
    from tests.conftest import chat_request

    _enable_confirmed_generation(client)
    uid = free_user.get("user_id")
    if uid is None:
        con = sqlite3.connect(tmp_db_path)
        uid = con.execute("SELECT id FROM users WHERE email LIKE 'test-free-user%'").fetchone()[0]
        con.close()
    oid = go.create(
        uid, "xlsx", "reproduce data from image",
        ask_content="[1 image(s) attached for visual context]\n"
                    "Give me an Excel reproduction of what you see in this image",
        images=["ZmFrZWpwZWc="])
    monkeypatch.setattr(dg, "interpret_offer_reply",
                        AsyncMock(return_value={"confirm": True, "format": "xlsx"}))

    r = client.post("/v1/chat", json=chat_request(
        prompt_mode="PostMeetingChat", call_type="query",
        metadata={"offer_id": oid, "generation_id": "gen-img-yes"},
        user_content="Yes",
    ), headers=free_user["headers"])
    assert r.status_code == 200
    assert r.headers["content-type"].startswith("text/event-stream")
    assert "event: generation_result" in r.text
    con = sqlite3.connect(tmp_db_path)
    row = con.execute("SELECT status FROM generations WHERE generation_id='gen-img-yes'").fetchone()
    con.close()
    assert row and row[0] == "done"


def test_typed_yes_image_guard_disarms_when_photo_missing(client, free_user, mock_provider,
                                                          tmp_db_path, monkeypatch):
    """Image-sourced ask, NO image anywhere (offer minted without images,
    reply send image-less): the guard disarms instead of generating —
    an image-blind build invents rather than fails (2026-07-19: fabricated
    sales sheet with placeholder names). The turn answers as normal chat
    with the re-attach steering, and must NOT re-offer."""
    import sqlite3
    from unittest.mock import AsyncMock
    import app.services.document_generation as dg
    from app.services import generation_offers as go
    from tests.conftest import chat_request

    _enable_confirmed_generation(client)
    uid = free_user.get("user_id")
    if uid is None:
        con = sqlite3.connect(tmp_db_path)
        uid = con.execute("SELECT id FROM users WHERE email LIKE 'test-free-user%'").fetchone()[0]
        con.close()
    oid = go.create(
        uid, "xlsx", "reproduce data from image",
        ask_content="[1 image(s) attached for visual context]\n"
                    "Give me an Excel reproduction of what you see in this image")
    monkeypatch.setattr(dg, "interpret_offer_reply",
                        AsyncMock(return_value={"confirm": True, "format": "xlsx"}))

    r = client.post("/v1/chat", json=chat_request(
        prompt_mode="PostMeetingChat", call_type="query",
        metadata={"offer_id": oid, "generation_id": "gen-img-guard"},
        user_content="Yes",
    ), headers=free_user["headers"])
    assert r.status_code == 200
    assert r.headers["content-type"].startswith("application/json")  # not SSE
    assert "feature_state" not in r.json()       # no re-offer loop
    con = sqlite3.connect(tmp_db_path)
    row = con.execute("SELECT 1 FROM generations WHERE generation_id='gen-img-guard'").fetchone()
    con.close()
    assert row is None                            # nothing generated


def test_chat_decline_is_a_normal_turn(client, free_user, mock_provider,
                                       tmp_db_path, monkeypatch):
    import sqlite3
    from unittest.mock import AsyncMock
    import app.services.document_generation as dg
    from app.services import generation_offers as go
    from tests.conftest import chat_request

    _enable_confirmed_generation(client)
    uid = free_user.get("user_id")
    if uid is None:
        con = sqlite3.connect(tmp_db_path)
        uid = con.execute("SELECT id FROM users WHERE email LIKE 'test-free-user%'").fetchone()[0]
        con.close()
    oid = go.create(uid, "docx", "for onboarding")
    monkeypatch.setattr(dg, "interpret_offer_reply",
                        AsyncMock(return_value={"confirm": False, "format": "docx"}))

    r = client.post("/v1/chat", json=chat_request(
        prompt_mode="ProjectChat", call_type="query",
        metadata={"offer_id": oid, "generation_id": "gen-chat-no"},
        user_content="no thanks, what time is the standup?",
    ), headers=free_user["headers"])
    assert r.status_code == 200
    assert r.headers["content-type"].startswith("application/json")
    assert "feature_state" not in r.json()      # mock classifier junk fails open
    con = sqlite3.connect(tmp_db_path)
    row = con.execute("SELECT 1 FROM generations WHERE generation_id='gen-chat-no'").fetchone()
    con.close()
    assert row is None                           # ids discarded on decline


def test_chat_confirm_resolution_splits_provider_prefixed_rows(client, free_user, mock_provider,
                                                               tmp_db_path, monkeypatch):
    """First live chat-confirm failed on 'anthropic/claude-sonnet-4-6': the
    lane re-resolution must split provider/model rows exactly like the
    first-pass resolution — the provider must NEVER see a vendor prefix."""
    import sqlite3
    from unittest.mock import AsyncMock
    import app.routers.chat as chat_mod
    import app.services.document_generation as dg
    from app.services import generation_offers as go
    from tests.conftest import chat_request

    _enable_confirmed_generation(client)
    uid = free_user.get("user_id")
    if uid is None:
        con = sqlite3.connect(tmp_db_path)
        uid = con.execute("SELECT id FROM users WHERE email LIKE 'test-free-user%'").fetchone()[0]
        con.close()
    oid = go.create(uid, "xlsx", "gantt for tasks")
    monkeypatch.setattr(dg, "interpret_offer_reply",
                        AsyncMock(return_value={"confirm": True, "format": "xlsx"}))
    monkeypatch.setattr(chat_mod, "_resolve_model_routing",
                        lambda *a, **k: "anthropic/claude-sonnet-4-6")

    r = client.post("/v1/chat", json=chat_request(
        prompt_mode="ProjectChat", call_type="query",
        metadata={"offer_id": oid, "generation_id": "gen-split-test"},
        user_content="Yeah",
    ), headers=free_user["headers"])
    assert r.status_code == 200
    assert "event: generation_result" in r.text          # NOT generation_error
    # the request the provider actually received: prefix split, never a "/"
    sent = mock_provider.await_args_list[-1].args[0]
    assert sent.provider == "anthropic"
    assert sent.model == "claude-sonnet-4-6"


def test_model_not_found_error_hides_the_catalog(client):
    import pytest as _pytest
    from fastapi import HTTPException
    router = client.app.state.provider_router
    with _pytest.raises(HTTPException) as ei:
        router.validate_model("anthropic", "claude-nonexistent-9")
    msg = ei.value.detail["message"]
    assert "Available" not in msg
    assert "claude" not in msg.lower()                   # no catalog ids leak


# --- Meeting Chat parity + classifier prefilter ---

def test_prefilter_vocabulary():
    from app.services.document_generation import looks_like_file_ask
    assert looks_like_file_ask("can you make a spreadsheet of the top 4")
    assert looks_like_file_ask("quiero un informe de la reunión")
    assert looks_like_file_ask("会議のレポートを作って")
    assert not looks_like_file_ask("what time is the standup tomorrow?")
    assert not looks_like_file_ask("who owns the auth fix?")


@pytest.mark.asyncio
async def test_classifier_skips_llm_when_prefilter_misses():
    from unittest.mock import AsyncMock, MagicMock
    from app.services.document_generation import classify_generation_intent
    router = MagicMock()
    router.route = AsyncMock()
    out = await classify_generation_intent(router, "what time is the standup?")
    assert out is None
    router.route.assert_not_awaited()          # no LLM call, no latency tax


def test_meeting_chat_file_ask_draws_offer_despite_stream_true(client, free_user,
                                                               mock_provider, monkeypatch):
    from unittest.mock import AsyncMock
    import app.services.document_generation as dg
    from tests.conftest import chat_request

    _enable_confirmed_generation(client)
    monkeypatch.setattr(dg, "classify_generation_intent", AsyncMock(
        return_value={"file_request": True, "format": "xlsx", "gist": "of action items"}))
    r = client.post("/v1/chat", json=chat_request(
        prompt_mode="PostMeetingChat", call_type="meeting_chat",
        stream=True, user_content="make a spreadsheet of the action items",
    ), headers=free_user["headers"])
    assert r.status_code == 200
    assert r.headers["content-type"].startswith("application/json")   # single JSON on the SSE request
    fs = r.json()["feature_state"]
    assert fs["state"] == "confirmation_required"
    assert fs["cta"]["details"]["offer_id"]


def test_meeting_chat_confirm_rides_generation_sse(client, free_user, mock_provider,
                                                   tmp_db_path, monkeypatch):
    import sqlite3
    from unittest.mock import AsyncMock
    import app.services.document_generation as dg
    from app.services import generation_offers as go
    from tests.conftest import chat_request

    _enable_confirmed_generation(client)
    uid = free_user.get("user_id")
    if uid is None:
        con = sqlite3.connect(tmp_db_path)
        uid = con.execute("SELECT id FROM users WHERE email LIKE 'test-free-user%'").fetchone()[0]
        con.close()
    oid = go.create(uid, "xlsx", "of action items")
    monkeypatch.setattr(dg, "interpret_offer_reply",
                        AsyncMock(return_value={"confirm": True, "format": "xlsx"}))
    r = client.post("/v1/chat", json=chat_request(
        prompt_mode="PostMeetingChat", call_type="meeting_chat", stream=True,
        metadata={"offer_id": oid, "generation_id": "gen-mc-parity"},
        user_content="yes please",
    ), headers=free_user["headers"])
    assert r.status_code == 200
    assert r.headers["content-type"].startswith("text/event-stream")
    assert "event: generation_result" in r.text


def test_isolate_reply_from_reinjected_attachment():
    from app.services.document_generation import _isolate_reply
    template = "RAG status: Red/Yellow/Green. " * 60          # ~1.8K of template
    assembled = f'--- Attached: "Asana report" ---\n{template}\n--- End ---\nCurrent question: Yes'
    assert _isolate_reply(assembled) == "Yes"                  # the actual reply, nothing else
    assembled2 = f'context stuff\nUser question: yes, the word doc please'
    assert _isolate_reply(assembled2) == "yes, the word doc please"
    # no marker -> plain tail (pre-injection behavior)
    assert _isolate_reply("just go ahead") == "just go ahead"


@pytest.mark.asyncio
async def test_interpreter_judges_reply_not_template(monkeypatch):
    from unittest.mock import AsyncMock, MagicMock
    from app.services.document_generation import interpret_offer_reply
    router = MagicMock()
    router.route = AsyncMock(return_value=MagicMock(text='{"confirm": true, "format": null}'))
    template = "Variance: [Variance]. Red/Yellow escalation. " * 50
    assembled = f'--- Attached: "Asana report" ---\n{template}\nCurrent question: Yes'
    out = await interpret_offer_reply(router, {"format": "docx", "gist": "x"}, assembled)
    sent = router.route.await_args.args[0].user_content
    assert "USER REPLY: Yes" in sent
    assert "Red/Yellow" not in sent               # template never reaches the judge
    assert out == {"confirm": True, "format": "docx"}


@pytest.mark.asyncio
async def test_interpreter_prefers_verbatim_reply_text():
    from unittest.mock import AsyncMock, MagicMock
    from app.services.document_generation import interpret_offer_reply
    router = MagicMock()
    router.route = AsyncMock(return_value=MagicMock(text='{"confirm": true, "format": null}'))
    # verbatim=True bypasses marker isolation entirely — SS sends the raw reply
    out = await interpret_offer_reply(
        router, {"format": "xlsx", "gist": "x"},
        "Yes", verbatim=True)
    sent = router.route.await_args.args[0].user_content
    assert sent.endswith("USER REPLY: Yes")
    assert out["confirm"] is True


# --- template registry (Smartsheet Gantt pilot) ---

_PLAN = {
    "project": "ABM IT Helpdesk",
    "tasks": [
        {"id": 1, "name": "Authentication / 401 Fix", "type": "phase", "parent_id": None,
         "owner": None, "status": "in_progress", "start": "2026-07-08", "end": "2026-07-15",
         "depends_on": []},
        {"id": 2, "name": "Share proxy logs", "type": "task", "parent_id": 1,
         "owner": "Sarah Park", "status": "complete", "start": "2026-07-08",
         "end": "2026-07-08", "depends_on": []},
        {"id": 3, "name": "Investigate 401s", "type": "task", "parent_id": 1,
         "owner": "Chirag Amin", "status": "blocked", "start": "2026-07-08",
         "end": "2026-07-14", "depends_on": [2]},
        {"id": 4, "name": "Verify auth-profile", "type": "task", "parent_id": 1,
         "owner": "Tom Lee", "status": "not_started", "start": "2026-07-09",
         "end": "2026-07-11", "depends_on": []},
        {"id": 5, "name": "Production release", "type": "milestone", "parent_id": 1,
         "owner": None, "status": "not_started", "start": "2026-07-15",
         "end": "2026-07-15", "depends_on": [3]},
    ],
}


def test_gantt_renderer_matches_reference_vocabulary():
    import datetime
    import io
    import openpyxl
    from app.services.doc_templates import render_gantt
    blob = render_gantt(_PLAN, today=datetime.date(2026, 7, 12))
    wb = openpyxl.load_workbook(io.BytesIO(blob))
    ws = wb["Gantt View"]
    assert ws.freeze_panes == "J4"
    owners = [str(c.value) for row in ws.iter_rows(min_col=9, max_col=9) for c in row if c.value]
    assert "Sarah Park" in owners and "Chirag Amin" in owners   # full names beside chips
    levels = {ws.row_dimensions[r].outline_level
              for r in ws.row_dimensions if ws.row_dimensions[r].outline_level}
    assert levels == {1, 2}                       # working collapse hierarchy
    texts = [str(c.value) for row in ws.iter_rows() for c in row if c.value]
    joined = " ".join(texts)
    assert "STATUS KEY" in joined
    assert "\U0001f3c1" in joined or "🏁" in joined  # milestone flag
    assert "⚑" in joined and "⚐" in joined          # at-risk flags both states
    assert "◆" in joined                             # milestone marker (formula-driven)
    # bars are ALSO characters (Numbers renders no conditional formatting —
    # live 2026-07-16): per-cell block-char formulas keyed to the date cells
    assert 'IF(AND(' in joined and "█" in joined
    # Predecessors column: dates-derived nomenclature (Scott 2026-07-15).
    # Task 3 shares task 2's start -> SS; milestone 5 follows 3's end -> FS.
    preds = [str(c.value) for row in ws.iter_rows(min_col=7, max_col=7)
             for c in row if c.value and re.fullmatch(r"[\d]+(FS|SS|FF)(, [\d]+(FS|SS|FF))*", str(c.value))]
    assert any(p.endswith("SS") for p in preds)
    assert any(p.endswith("FS") for p in preds)
    # LIVE GRID: bars/today/weekend/risk are conditional-formatting rules
    # over the hidden real-date axis row, so user date edits redraw them
    assert ws.row_dimensions[1].hidden
    axis = [c.value for c in ws[1] if c.value is not None]
    assert all(isinstance(v, datetime.datetime) for v in axis)
    rules = [r for rng in ws.conditional_formatting for r in rng.rules]
    formulas = " | ".join(f for r in rules for f in (r.formula or []))
    assert "TODAY()" in formulas                      # live today + risk logic
    assert "WEEKDAY(" in formulas                     # live weekend shading
    assert ">=$E" in formulas and "<=$F" in formulas  # bars keyed to date cells
    # reference palette, exact — now carried by the CF fills + static bands
    cf_hex = {r.dxf.fill.fgColor.rgb for r in rules
              if r.dxf is not None and r.dxf.fill is not None
              and isinstance(r.dxf.fill.fgColor.rgb, str)}
    static_hex = {c.fill.fgColor.rgb for row in ws.iter_rows() for c in row
                  if c.fill and isinstance(c.fill.fgColor.rgb, str)}
    fills = cf_hex | static_hex
    for hex6 in ("FFA8B9C9", "FF6E7B8A", "FF3D4653", "FFF3F3F3", "FFFFF6DE", "FFE0341E"):
        assert hex6 in fills, hex6
    # start/end are real dates (formulas compare against them)
    d_cells = [c.value for row in ws.iter_rows(min_col=5, max_col=6) for c in row
               if isinstance(c.value, datetime.datetime)]
    assert d_cells, "start/end must be date-typed for the live grid"
    # status dropdown present (drives dot recolor + risk rule)
    assert ws.data_validations.dataValidation
    # determinism: same plan, same bytes
    assert render_gantt(_PLAN, today=datetime.date(2026, 7, 12)) == blob


def test_dep_code_derivation_unit():
    from app.services.doc_templates import _dep_code
    a = {"start": "2026-07-01", "end": "2026-07-05"}
    assert _dep_code(a, {"start": "2026-07-06", "end": "2026-07-09"}) == "FS"
    assert _dep_code(a, {"start": "2026-07-01", "end": "2026-07-09"}) == "SS"
    assert _dep_code(a, {"start": "2026-07-03", "end": "2026-07-05"}) == "FF"
    # ambiguity defaults to FS (overlap without aligned edges)
    assert _dep_code(a, {"start": "2026-07-03", "end": "2026-07-08"}) == "FS"


def test_template_match_and_parse():
    from app.services.doc_templates import match_template, parse_extraction
    assert match_template("can you build a gantt chart of our plan?") == "gantt_smartsheet"
    assert match_template("make me a spreadsheet of insults") is None
    assert parse_extraction('Sure: {"project": "X", "tasks": []} done') == {"project": "X", "tasks": []}


def test_template_offer_intercepts_and_confirm_renders(client, free_user, mock_provider,
                                                       tmp_db_path, monkeypatch):
    import json as _json
    import sqlite3
    from unittest.mock import AsyncMock
    import app.services.document_generation as dg
    from app.services import generation_offers as go
    from tests.conftest import chat_request

    _enable_confirmed_generation(client)
    monkeypatch.setattr(dg, "classify_generation_intent", AsyncMock(
        return_value={"file_request": True, "format": "xlsx", "gist": "of our project plan"}))
    r = client.post("/v1/chat", json=chat_request(
        prompt_mode="ProjectChat", call_type="query",
        user_content="build a gantt chart of our project plan",
    ), headers=free_user["headers"])
    cta = r.json()["feature_state"]["cta"]
    assert cta["details"]["template_id"] == "gantt_smartsheet"
    assert "polished" in cta["text"] and "custom" in cta["text"]   # open door
    oid = cta["details"]["offer_id"]

    # confirm: extraction turn (mock provider returns the plan JSON), then
    # the deterministic renderer stages a real xlsx
    uid = free_user.get("user_id")
    if uid is None:
        con = sqlite3.connect(tmp_db_path)
        uid = con.execute("SELECT id FROM users WHERE email LIKE 'test-free-user%'").fetchone()[0]
        con.close()
    monkeypatch.setattr(dg, "interpret_offer_reply", AsyncMock(
        return_value={"confirm": True, "format": "xlsx"}))
    mock_provider.canned_response.text = _json.dumps(_PLAN)
    mock_provider.return_value = mock_provider.canned_response
    r2 = client.post("/v1/chat", json=chat_request(
        prompt_mode="ProjectChat", call_type="query",
        metadata={"offer_id": oid, "generation_id": "gen-tmpl-1"},
        user_content="yes the polished one",
    ), headers=free_user["headers"])
    assert r2.status_code == 200
    assert r2.headers["content-type"].startswith("text/event-stream")
    result = _json.loads(r2.text.split("event: generation_result\ndata: ")[1].split("\n")[0])
    assert re.fullmatch(r"([A-Za-z0-9_]+_)?Gantt_\d{6}\.xlsx",
                        result["generated_files"][0]["name"])
    assert "Built your xlsx" in result["text"]
    # extraction turn was NOT a sandbox turn: no container arming happened
    sent = mock_provider.await_args_list[-1].args[0]
    assert sent.generation is False
    assert "JSON" in sent.system_prompt          # extraction prompt replaced client prompt


def test_template_match_survives_long_prompts_with_leading_keyword():
    from app.services.doc_templates import match_template
    long_spec = "Build me a project Gantt chart as an Excel file. " + ("Timeline details and colors. " * 200)
    assert len(long_spec) > 2000                      # keyword far outside any tail window
    assert match_template(long_spec) == "gantt_smartsheet"


def test_confirmed_turn_runs_on_originating_ask_content(client, free_user, mock_provider,
                                                        tmp_db_path, monkeypatch):
    """First live template run: the reply send carried 410 chars of chat
    history and the extraction asked the USER for the plan. The confirmed
    turn must run against the originating ask's content instead."""
    import json as _json
    import sqlite3
    from unittest.mock import AsyncMock
    import app.services.document_generation as dg
    from app.services import generation_offers as go
    from tests.conftest import chat_request

    _enable_confirmed_generation(client)
    meeting_context = ("MEETING CONTENT: auth phase Jul 8-15 owned by Chirag; "
                       "release milestone Jul 15. " * 20)
    monkeypatch.setattr(dg, "classify_generation_intent", AsyncMock(
        return_value={"file_request": True, "format": "xlsx", "gist": "plan"}))
    r = client.post("/v1/chat", json=chat_request(
        prompt_mode="ProjectChat", call_type="query",
        user_content=meeting_context + " build a gantt chart of this",
    ), headers=free_user["headers"])
    oid = r.json()["feature_state"]["cta"]["details"]["offer_id"]

    monkeypatch.setattr(dg, "interpret_offer_reply", AsyncMock(
        return_value={"confirm": True, "format": "xlsx"}))
    mock_provider.canned_response.text = _json.dumps(_PLAN)
    mock_provider.return_value = mock_provider.canned_response
    r2 = client.post("/v1/chat", json=chat_request(
        prompt_mode="ProjectChat", call_type="query",
        metadata={"offer_id": oid, "generation_id": "gen-ctx-1",
                  "reply_text": "yes"},
        user_content="Previous conversation: Q/A only. Current question: yes",
    ), headers=free_user["headers"])
    assert r2.status_code == 200
    sent = mock_provider.await_args_list[-1].args[0]
    assert "MEETING CONTENT" in sent.user_content            # originating context restored
    assert "confirmed the file build with: yes" in sent.user_content
    assert "Previous conversation: Q/A only" not in sent.user_content
    # the client's system prompt (where Project Chat carries meeting
    # summaries) survives, with the extraction directive appended
    assert sent.system_prompt.startswith("You are a helpful assistant.")
    assert "FILE BUILD OVERRIDE" in sent.system_prompt
    assert "JSON ONLY" in sent.system_prompt


def test_template_lane_stamps_generation_metering(client, free_user, mock_provider,
                                                  tmp_db_path, monkeypatch):
    import json as _json
    import sqlite3
    from unittest.mock import AsyncMock
    import app.services.document_generation as dg
    from app.services import generation_offers as go
    from tests.conftest import chat_request

    _enable_confirmed_generation(client)
    uid = free_user.get("user_id")
    if uid is None:
        con = sqlite3.connect(tmp_db_path)
        uid = con.execute("SELECT id FROM users WHERE email LIKE 'test-free-user%'").fetchone()[0]
        con.close()
    oid = go.create(uid, "xlsx", "plan", template_id="gantt_smartsheet",
                    ask_content="meetings say things")
    monkeypatch.setattr(dg, "interpret_offer_reply", AsyncMock(
        return_value={"confirm": True, "format": "xlsx"}))
    mock_provider.canned_response.text = _json.dumps(_PLAN)
    mock_provider.return_value = mock_provider.canned_response
    r = client.post("/v1/chat", json=chat_request(
        prompt_mode="ProjectChat", call_type="query",
        metadata={"offer_id": oid, "generation_id": "gen-meter-1", "reply_text": "yes"},
        user_content="yes",
    ), headers=free_user["headers"])
    assert r.status_code == 200
    con = sqlite3.connect(tmp_db_path)
    con.row_factory = sqlite3.Row
    row = con.execute("SELECT metadata FROM usage_log WHERE user_id = ? "
                      "ORDER BY request_timestamp DESC LIMIT 1", (uid,)).fetchone()
    con.close()
    md = _json.loads(row["metadata"])
    assert md.get("generated") == {"count": 1, "bytes": md["generated"]["bytes"]}
    assert md["generated"]["bytes"] > 1000


def test_report_stoplight_is_calibrated():
    # TR field report: same conversation, different verdicts. Temperature
    # pinned + evidence-anchored bands (same closed-enum lesson as
    # compare-reality's delta field).
    src = open("app/routers/reports.py").read()
    import re
    block = src[src.index("chat_request = ChatRequest("):]
    assert "temperature=0.2" in block[:1200]
    prompt = open("app/services/meeting_report.py").read()
    assert "EVIDENCE-ANCHORED" in prompt
    assert "no agreed path forward" in prompt          # red criteria observable
    assert "Borderline rule" in prompt                 # ties resolve deterministically


def test_explicit_file_asks_are_a_guaranteed_catch():
    from app.services.document_generation import explicit_file_ask
    for ask, fmt in (("make me a file of the risks", None),
                     ("please generate a spreadsheet from this", "xlsx"),
                     ("build the docx we discussed", "docx"),
                     ("can you create a PowerPoint for the exec review", "pptx")):
        out = explicit_file_ask("meeting context blah. " * 50 + ask)
        assert out is not None and out["file_request"] is True, ask
        assert out["format"] == fmt, ask
    assert explicit_file_ask("what did we decide about the file server?") is None


def test_explicit_catch_offers_even_when_classifier_is_down(client, free_user,
                                                            mock_provider, monkeypatch):
    from unittest.mock import AsyncMock
    import app.services.document_generation as dg
    from tests.conftest import chat_request
    _enable_confirmed_generation(client)
    monkeypatch.setattr(dg, "classify_generation_intent",
                        AsyncMock(side_effect=RuntimeError("classifier down")))
    r = client.post("/v1/chat", json=chat_request(
        prompt_mode="ProjectChat", call_type="query",
        user_content="generate a spreadsheet of our action items",
    ), headers=free_user["headers"])
    assert r.json()["feature_state"]["state"] == "confirmation_required"


def test_soft_vocabulary_gets_teaser_not_offer(client, free_user, mock_provider, monkeypatch):
    from unittest.mock import AsyncMock
    import app.services.document_generation as dg
    from tests.conftest import chat_request
    _enable_confirmed_generation(client)
    monkeypatch.setattr(dg, "classify_generation_intent", AsyncMock(
        return_value={"file_request": False, "format": None, "gist": ""}))
    r = client.post("/v1/chat", json=chat_request(
        prompt_mode="ProjectChat", call_type="query",
        user_content="what did the report say about the deadline?",
    ), headers=free_user["headers"])
    body = r.json()
    assert body["text"]                                        # normal answer
    fs = body["feature_state"]
    assert fs["state"] == "available"
    assert fs["cta"]["kind"] == "generation_teaser"
    assert fs["cta"]["text"] == "Want this as a real downloadable file?"


def test_report_chat_request_carries_request_id_metadata():
    src = open("app/routers/reports.py").read()
    block = src[src.index("chat_request = ChatRequest("):]
    assert '"request_id": getattr(request.state, "request_id"' in block[:800]


def test_reports_route_carries_tr_budget_pre_gate():
    src = open("app/routers/reports.py").read()
    i = src.index("techrehearsal")
    block = src[i:i+900]
    assert "would_exceed_tr_budget" in block
    assert "429" in block and "allocation_exhausted" in block   # TR's exact shape


def test_generation_survives_client_disconnect(client, free_user, mock_provider,
                                               tmp_db_path, monkeypatch):
    """Kill-test hardening: the client vanishing mid-stream must not stop
    the turn — the rescue row and artifact are the whole point."""
    import json as _json
    import sqlite3
    import time
    from unittest.mock import AsyncMock
    import app.services.document_generation as dg
    from app.services import generation_offers as go
    from tests.conftest import chat_request

    _enable_confirmed_generation(client)
    uid = free_user.get("user_id")
    if uid is None:
        con = sqlite3.connect(tmp_db_path)
        uid = con.execute("SELECT id FROM users WHERE email LIKE 'test-free-user%'").fetchone()[0]
        con.close()
    oid = go.create(uid, "xlsx", "plan", template_id="gantt_smartsheet",
                    ask_content="ctx")
    monkeypatch.setattr(dg, "interpret_offer_reply", AsyncMock(
        return_value={"confirm": True, "format": "xlsx"}))
    mock_provider.canned_response.text = _json.dumps(_PLAN)
    mock_provider.return_value = mock_provider.canned_response

    # open the stream and abandon it after the FIRST event (simulated kill)
    with client.stream("POST", "/v1/chat", json=chat_request(
            prompt_mode="ProjectChat", call_type="query",
            metadata={"offer_id": oid, "generation_id": "gen-kill-1", "reply_text": "yes"},
            user_content="yes"), headers=free_user["headers"]) as r:
        for _ in r.iter_lines():
            break                                   # got generation_started; die

    # the turn must still complete and write its rescue row
    deadline = time.time() + 15
    row = None
    while time.time() < deadline and row is None:
        con = sqlite3.connect(tmp_db_path)
        row = con.execute("SELECT status FROM generations WHERE generation_id='gen-kill-1'").fetchone()
        con.close()
        if row is None:
            time.sleep(0.5)
    assert row is not None and row[0] == "done"


def test_intent_checks_ignore_persistent_injection_blocks():
    """Conversation-scoped attachments ride every turn — the injection
    text must not re-trigger intent machinery on unrelated follow-ups."""
    from app.services.document_generation import _question_portion, explicit_file_ask, looks_like_file_ask
    turn = ('--- Attached: "plan.xlsx" ---\nspreadsheet of files and documents\n'
            '--- End ---\nCurrent question: who owns the auth fix?')
    assert _question_portion(turn) == "who owns the auth fix?"
    assert not looks_like_file_ask(_question_portion(turn))
    assert explicit_file_ask(turn) is None            # despite "spreadsheet" in the block
    turn2 = turn.replace("who owns the auth fix?", "now generate a spreadsheet of it")
    assert explicit_file_ask(turn2)["file_request"] is True


def test_template_confirm_on_streaming_surface_rides_generation_transport(
        client, free_user, mock_provider, tmp_db_path, monkeypatch):
    """Live bug 2026-07-13 18:54Z: a confirmed template turn on meeting
    chat (stream=true) token-streamed the raw extraction JSON into the
    chat and never ran the renderer. Template-armed turns must divert to
    the generation transport on streaming surfaces, exactly like
    sandbox-armed ones."""
    import json as _json
    import sqlite3
    from unittest.mock import AsyncMock
    import app.services.document_generation as dg
    from app.services import generation_offers as go
    from tests.conftest import chat_request

    _enable_confirmed_generation(client)
    uid = free_user.get("user_id")
    if uid is None:
        con = sqlite3.connect(tmp_db_path)
        uid = con.execute("SELECT id FROM users WHERE email LIKE 'test-free-user%'").fetchone()[0]
        con.close()
    oid = go.create(uid, "xlsx", "gantt of open items",
                    template_id="gantt_smartsheet",
                    ask_content="Build a nice Gantt chart of the open items")
    monkeypatch.setattr(dg, "interpret_offer_reply", AsyncMock(
        return_value={"confirm": True, "format": "xlsx"}))
    mock_provider.canned_response.text = _json.dumps(_PLAN)
    mock_provider.return_value = mock_provider.canned_response

    r = client.post("/v1/chat", json=chat_request(
        prompt_mode="PostMeetingChat", call_type="meeting_chat", stream=True,
        metadata={"offer_id": oid, "generation_id": "gen-tmpl-mc-1",
                  "reply_text": "Yes"},
        user_content="Yes",
    ), headers=free_user["headers"])
    assert r.status_code == 200
    # the failure mode was a token stream of raw JSON; the contract is the
    # generation transport with a rendered file and friendly text
    assert r.headers["content-type"].startswith("text/event-stream")
    assert "event: generation_started" in r.text
    result = _json.loads(
        r.text.split("event: generation_result\ndata: ")[1].split("\n")[0])
    assert re.fullmatch(r"([A-Za-z0-9_]+_)?Gantt_\d{6}\.xlsx",
                        result["generated_files"][0]["name"])
    assert "Built your xlsx" in result["text"]
    assert not result["text"].lstrip().startswith("{")

    con = sqlite3.connect(tmp_db_path)
    row = con.execute("SELECT status FROM generations "
                      "WHERE generation_id='gen-tmpl-mc-1'").fetchone()
    con.close()
    assert row and row[0] == "done"


def test_parse_extraction_recovers_plan_from_narration_json_and_html():
    """Live 2026-07-13 19:16Z: the extraction turn narrated, emitted the
    plan JSON, then a full HTML page. The old first-{-to-last-} slice
    spanned into the HTML's CSS braces and failed a turn that carried a
    valid plan."""
    import json as _json
    from app.services.doc_templates import parse_extraction

    disobedient = (
        "I'll create a Smartsheet-style Gantt chart based on the meeting "
        "activities. Let me first extract the project plan, then render "
        "it visually.\n\n"
        + _json.dumps(_PLAN)
        + "\n\nNow here's the Smartsheet-style Gantt chart:\n\n"
        "<!DOCTYPE html>\n<html lang=\"en\">\n<head>\n<style>\n"
        "* { box-sizing: border-box; margin: 0; padding: 0; }\n"
        ".ss-header { background: linear-gradient(135deg, #0a2342 0%, "
        "#123); color: #fff; padding: 14px 24px; }\n"
        "</style>\n</head>\n<body></body>\n</html>"
    )
    assert parse_extraction(disobedient) == _PLAN
    # a stray small object in the narration must not shadow the plan
    assert parse_extraction('{"note": "extracting"} ' + _json.dumps(_PLAN)) == _PLAN
    # clean output still parses; garbage still raises
    assert parse_extraction(_json.dumps(_PLAN)) == _PLAN
    import pytest as _pytest
    with _pytest.raises(Exception):
        parse_extraction("no json here at all")


def test_template_render_survives_disobedient_extraction_e2e(
        client, free_user, mock_provider, tmp_db_path, monkeypatch):
    """Same flow as the 19:16Z failure, e2e: model narrates + JSON + HTML,
    the renderer must still produce the file and friendly text."""
    import json as _json
    import sqlite3
    from unittest.mock import AsyncMock
    import app.services.document_generation as dg
    from app.services import generation_offers as go
    from tests.conftest import chat_request

    _enable_confirmed_generation(client)
    uid = free_user.get("user_id")
    if uid is None:
        con = sqlite3.connect(tmp_db_path)
        uid = con.execute("SELECT id FROM users WHERE email LIKE 'test-free-user%'").fetchone()[0]
        con.close()
    oid = go.create(uid, "xlsx", "gantt of open items",
                    template_id="gantt_smartsheet",
                    ask_content="Build a Gantt chart of open items")
    monkeypatch.setattr(dg, "interpret_offer_reply", AsyncMock(
        return_value={"confirm": True, "format": "xlsx"}))
    mock_provider.canned_response.text = (
        "Let me extract then render.\n" + _json.dumps(_PLAN)
        + "\n<!DOCTYPE html><html><head><style>body { margin: 0; }"
          "</style></head><body></body></html>")
    mock_provider.return_value = mock_provider.canned_response

    r = client.post("/v1/chat", json=chat_request(
        prompt_mode="PostMeetingChat", call_type="meeting_chat", stream=True,
        metadata={"offer_id": oid, "generation_id": "gen-tmpl-html-1",
                  "reply_text": "Yes"},
        user_content="Yes",
    ), headers=free_user["headers"])
    assert r.status_code == 200
    assert r.headers["content-type"].startswith("text/event-stream")
    result = _json.loads(
        r.text.split("event: generation_result\ndata: ")[1].split("\n")[0])
    assert re.fullmatch(r"([A-Za-z0-9_]+_)?Gantt_\d{6}\.xlsx",
                        result["generated_files"][0]["name"])
    assert "Built your xlsx" in result["text"]
    assert "<!DOCTYPE" not in result["text"]


def test_bare_followup_in_file_heavy_conversation_gets_no_teaser(
        client, free_user, mock_provider):
    """Live 2026-07-14: 'Test' sent into a post-generation Project Chat
    drew the teaser — the hint check scanned the whole assembled tail
    (gantt/file/xlsx vocabulary everywhere) instead of the question
    portion (#420's rule). A bare follow-up must be a plain answer."""
    from tests.conftest import chat_request
    _enable_confirmed_generation(client)
    assembled = (
        "Previous conversation in this chat: "
        "Q: Build a Gantt chart from these meetings A: Built your xlsx — "
        "38 tasks and milestones (Project_Gantt.xlsx, a native Excel "
        "file). Download the spreadsheet and open the report.\n\n"
        "Current question: Test"
    )
    r = client.post("/v1/chat", json=chat_request(
        prompt_mode="ProjectChat", call_type="query",
        user_content=assembled,
    ), headers=free_user["headers"])
    body = r.json()
    assert body["text"]
    fs = body.get("feature_state") or {}
    cta = fs.get("cta") or {}
    assert cta.get("kind") != "generation_teaser"


def test_file_vocabulary_in_the_question_itself_still_teases(
        client, free_user, mock_provider, monkeypatch):
    from unittest.mock import AsyncMock
    import app.services.document_generation as dg
    from tests.conftest import chat_request
    _enable_confirmed_generation(client)
    monkeypatch.setattr(dg, "classify_generation_intent", AsyncMock(
        return_value={"file_request": False, "format": None, "gist": ""}))
    r = client.post("/v1/chat", json=chat_request(
        prompt_mode="ProjectChat", call_type="query",
        user_content=("Previous conversation in this chat: Q: hello A: hi\n\n"
                      "Current question: what did the report say about the deadline?"),
    ), headers=free_user["headers"])
    fs = r.json()["feature_state"]
    assert fs["cta"]["kind"] == "generation_teaser"


def test_capable_unarmed_turn_carries_file_capability_line(
        client, free_user, mock_provider):
    """Live 2026-07-14: on a normal turn the model told a generation-
    capable user it 'doesn't have the ability to generate files'. Gate-
    passing UNARMED turns now append a server-side capability line (gate
    state is per-turn server knowledge; a static client line would lie
    to Free/BYOK users)."""
    from tests.conftest import chat_request
    _enable_confirmed_generation(client)
    r = client.post("/v1/chat", json=chat_request(
        prompt_mode="ProjectChat", call_type="query",
        system_prompt="You are a helpful assistant.",
        user_content="Current question: summarize the last meeting",
    ), headers=free_user["headers"])
    assert r.status_code == 200
    sent = mock_provider.await_args_list[-1].args[0]
    assert "FILE CAPABILITY" in sent.system_prompt
    assert sent.system_prompt.startswith("You are a helpful assistant.")


def test_gate_failing_turn_gets_no_capability_line(client, free_user, mock_provider):
    from tests.conftest import chat_request
    # generation NOT enabled -> gate fails -> stock prompt untouched
    docs = client.app.state.remote_configs["client-config"].setdefault("documents", {})
    docs["generation"] = {"enabled": False}
    r = client.post("/v1/chat", json=chat_request(
        prompt_mode="ProjectChat", call_type="query",
        user_content="Current question: summarize the last meeting",
    ), headers=free_user["headers"])
    assert r.status_code == 200
    sent = mock_provider.await_args_list[-1].args[0]
    assert "FILE CAPABILITY" not in (sent.system_prompt or "")


def test_armed_template_turn_gets_extraction_not_capability_line(
        client, free_user, mock_provider, tmp_db_path, monkeypatch):
    import json as _json
    import sqlite3
    from unittest.mock import AsyncMock
    import app.services.document_generation as dg
    from app.services import generation_offers as go
    from tests.conftest import chat_request

    _enable_confirmed_generation(client)
    uid = free_user.get("user_id")
    if uid is None:
        con = sqlite3.connect(tmp_db_path)
        uid = con.execute("SELECT id FROM users WHERE email LIKE 'test-free-user%'").fetchone()[0]
        con.close()
    oid = go.create(uid, "xlsx", "gantt", template_id="gantt_smartsheet",
                    ask_content="Build a gantt chart")
    monkeypatch.setattr(dg, "interpret_offer_reply", AsyncMock(
        return_value={"confirm": True, "format": "xlsx"}))
    mock_provider.canned_response.text = _json.dumps(_PLAN)
    mock_provider.return_value = mock_provider.canned_response
    r = client.post("/v1/chat", json=chat_request(
        prompt_mode="ProjectChat", call_type="query",
        metadata={"offer_id": oid, "generation_id": "gen-cap-t1"},
        user_content="yes",
    ), headers=free_user["headers"])
    assert r.status_code == 200
    sent = mock_provider.await_args_list[-1].args[0]
    assert "FILE BUILD OVERRIDE" in sent.system_prompt or "JSON" in sent.system_prompt
    assert "FILE CAPABILITY" not in sent.system_prompt


def test_teaser_mints_offer_and_typed_yes_arms(
        client, free_user, mock_provider, tmp_db_path, monkeypatch):
    """Joint call with SS 2026-07-14: teasers mint an offer so a TYPED
    yes rides the same echo lane as real offers (the pill tap keeps the
    generation_confirmed resend)."""
    import json as _json
    import sqlite3
    from unittest.mock import AsyncMock
    import app.services.document_generation as dg
    from tests.conftest import chat_request

    _enable_confirmed_generation(client)
    monkeypatch.setattr(dg, "classify_generation_intent", AsyncMock(
        return_value={"file_request": False, "format": None, "gist": ""}))
    r = client.post("/v1/chat", json=chat_request(
        prompt_mode="ProjectChat", call_type="query",
        user_content="Current question: what did the report say about the gantt chart?",
    ), headers=free_user["headers"])
    cta = r.json()["feature_state"]["cta"]
    assert cta["kind"] == "generation_teaser"
    oid = cta["details"]["offer_id"]
    assert oid

    monkeypatch.setattr(dg, "interpret_offer_reply", AsyncMock(
        return_value={"confirm": True, "format": "xlsx"}))
    mock_provider.canned_response.text = _json.dumps(_PLAN)
    mock_provider.return_value = mock_provider.canned_response
    r2 = client.post("/v1/chat", json=chat_request(
        prompt_mode="ProjectChat", call_type="query",
        metadata={"offer_id": oid, "generation_id": "gen-teaser-yes-1",
                  "reply_text": "Yes"},
        user_content="Yes",
    ), headers=free_user["headers"])
    assert r2.status_code == 200
    assert r2.headers["content-type"].startswith("text/event-stream")
    assert "event: generation_result" in r2.text
    con = sqlite3.connect(tmp_db_path)
    row = con.execute("SELECT status FROM generations "
                      "WHERE generation_id='gen-teaser-yes-1'").fetchone()
    con.close()
    assert row and row[0] == "done"


def test_pill_tap_at_teaser_inherits_the_minted_offer(
        client, free_user, mock_provider, tmp_db_path, monkeypatch):
    """SS's pill tap sends generation_confirmed AND the offer_id echo
    together (their 2026-07-14 fix). Confirmed used to short-circuit the
    reply-interpret block, so the echoed offer was never spent and its
    template_id/ask_content never inherited — a tap at a template-matched
    teaser rode the sandbox lane blind while a typed yes at the same
    teaser rode the template lane with full context."""
    import json as _json
    import sqlite3
    from unittest.mock import AsyncMock
    import app.services.document_generation as dg
    from tests.conftest import chat_request

    _enable_confirmed_generation(client)
    monkeypatch.setattr(dg, "classify_generation_intent", AsyncMock(
        return_value={"file_request": False, "format": None, "gist": ""}))
    meeting_context = ("MEETING CONTENT: auth phase Jul 8-15 owned by Chirag; "
                       "release milestone Jul 15. " * 20)
    r = client.post("/v1/chat", json=chat_request(
        prompt_mode="ProjectChat", call_type="query",
        user_content=meeting_context
        + " Current question: what did the report say about the gantt chart?",
    ), headers=free_user["headers"])
    cta = r.json()["feature_state"]["cta"]
    assert cta["kind"] == "generation_teaser"
    oid = cta["details"]["offer_id"]
    assert oid

    # the pill tap is consent already — the reply judge must NOT run
    _interpret = AsyncMock(side_effect=AssertionError(
        "interpret_offer_reply must not run on a confirmed turn"))
    monkeypatch.setattr(dg, "interpret_offer_reply", _interpret)
    mock_provider.canned_response.text = _json.dumps(_PLAN)
    mock_provider.return_value = mock_provider.canned_response
    r2 = client.post("/v1/chat", json=chat_request(
        prompt_mode="ProjectChat", call_type="query",
        metadata={"offer_id": oid, "generation_id": "gen-teaser-tap-1",
                  "generation_confirmed": True},
        user_content="Previous conversation: Q/A only. Current question: "
                     "what did the report say about the gantt chart?",
    ), headers=free_user["headers"])
    assert r2.status_code == 200
    assert r2.headers["content-type"].startswith("text/event-stream")
    result = _json.loads(r2.text.split(
        "event: generation_result\ndata: ")[1].split("\n")[0])
    # template lane, not sandbox: the registry renderer drew the file
    assert re.fullmatch(r"([A-Za-z0-9_]+_)?Gantt_\d{6}\.xlsx",
                        result["generated_files"][0]["name"])
    sent = mock_provider.await_args_list[-1].args[0]
    assert sent.generation is False
    assert "FILE BUILD OVERRIDE" in sent.system_prompt
    # originating ask inherited from the offer, decoy history dropped
    assert "MEETING CONTENT" in sent.user_content
    assert "confirmed the file build" in sent.user_content
    assert "Previous conversation: Q/A only" not in sent.user_content
    con = sqlite3.connect(tmp_db_path)
    row = con.execute("SELECT status FROM generations "
                      "WHERE generation_id='gen-teaser-tap-1'").fetchone()
    con.close()
    assert row and row[0] == "done"


def test_below_tier_explicit_file_ask_is_a_plain_chat_turn(
        client, free_user, mock_provider, monkeypatch):
    """Production gate shape (enabled, min_tier=pro) with a FREE user
    explicitly asking for a file. The tier check is unit-pinned in
    test_gate_matrix; this pins the e2e cell through /v1/chat: no offer,
    no teaser, no classifier spend, no capability line, no arming — a
    completely normal chat turn."""
    from unittest.mock import AsyncMock
    import app.services.document_generation as dg
    from tests.conftest import chat_request

    docs = client.app.state.remote_configs["client-config"].setdefault(
        "documents", {})
    docs["generation"] = {
        "enabled": True, "min_tier": "pro",
        "confirmation": {"enabled": True, "expected_seconds": 150}}
    monkeypatch.setattr(dg, "classify_generation_intent", AsyncMock(
        side_effect=AssertionError(
            "intent classifier must not run for a below-tier user")))
    r = client.post("/v1/chat", json=chat_request(
        prompt_mode="ProjectChat", call_type="query",
        user_content="Current question: create an excel file of our "
                     "project plan with all the milestones",
    ), headers=free_user["headers"])
    assert r.status_code == 200
    body = r.json()
    cta = (body.get("feature_state") or {}).get("cta") or {}
    assert cta.get("kind") not in ("generation_offer", "generation_teaser")
    assert "generated_files" not in body
    sent = mock_provider.await_args_list[-1].args[0]
    assert sent.generation is False
    assert "FILE CAPABILITY" not in (sent.system_prompt or "")
    # upsell not opted in (defaults ship dark) -> reply text untouched
    assert body["text"] == "Test response from mock provider."


# --- below-tier upsell line (Scott 2026-07-14) ---

def _enable_generation_with_upsell(client, min_tier="pro", upsell=None):
    docs = client.app.state.remote_configs["client-config"].setdefault(
        "documents", {})
    docs["generation"] = {
        "enabled": True, "min_tier": min_tier,
        "confirmation": {"enabled": True, "expected_seconds": 150},
        "upsell": upsell if upsell is not None else {"enabled": True},
    }


def test_tier_shortfall_matrix():
    from app.services.document_generation import generation_tier_shortfall

    def _short(**over):
        kw = dict(
            remote_configs={"client-config": {"documents": {"generation": {
                "enabled": True, "min_tier": "pro"}}}},
            tier_name="free", managed_routing=True, provider="anthropic",
            prompt_mode="ProjectChat",
        )
        kw.update(over)
        return generation_tier_shortfall(**kw)

    assert _short() == "pro"
    assert _short(tier_name="plus") == "pro"
    assert _short(prompt_mode="PostMeetingChat") == "pro"
    # at or above min_tier -> no shortfall (the real gate passes)
    assert _short(tier_name="pro") is None
    # unranked tiers (admin) fail the real gate's tier check today, but
    # they are never a shortfall — upselling an admin would be wrong
    assert _short(tier_name="admin") is None
    # any mechanical gate failure is NOT a tier shortfall
    assert _short(prompt_mode=None) is None
    assert _short(prompt_mode="InterviewScorecard") is None
    assert _short(managed_routing=False) is None
    assert _short(provider="openrouter") is None
    assert _short(remote_configs={"client-config": {"documents": {
        "generation": {"enabled": False, "min_tier": "pro"}}}}) is None


def test_below_tier_file_ask_gets_tier_aware_upsell_line(
        client, free_user, mock_provider, monkeypatch):
    """The served line leads the reply, {tier} resolves to the served
    min_tier's display name, the coherence line steers the model, and the
    classifier never runs (deterministic detection only)."""
    from unittest.mock import AsyncMock
    import app.services.document_generation as dg
    from tests.conftest import chat_request

    _enable_generation_with_upsell(client)
    monkeypatch.setattr(dg, "classify_generation_intent", AsyncMock(
        side_effect=AssertionError(
            "classifier must not run for a below-tier user")))
    r = client.post("/v1/chat", json=chat_request(
        prompt_mode="ProjectChat", call_type="query",
        user_content="Current question: create an excel file of our plan",
    ), headers=free_user["headers"])
    assert r.status_code == 200
    body = r.json()
    assert body["text"].startswith(
        "If you were a Pro subscriber, I could generate a Word or Excel "
        "file for you.\n\n")
    assert body["text"].endswith("Test response from mock provider.")
    sent = mock_provider.await_args_list[-1].args[0]
    assert "FILE UPSELL CONTEXT" in sent.system_prompt
    assert "Pro plan" in sent.system_prompt
    assert sent.generation is False
    cta = (body.get("feature_state") or {}).get("cta") or {}
    assert cta.get("kind") not in ("generation_offer", "generation_teaser")


def test_upsell_tier_name_follows_served_min_tier(
        client, free_user, mock_provider):
    """Scott's core requirement: never assume the feature is Pro. Move
    min_tier to plus in served config and the line follows."""
    from tests.conftest import chat_request

    _enable_generation_with_upsell(client, min_tier="plus")
    r = client.post("/v1/chat", json=chat_request(
        prompt_mode="ProjectChat", call_type="query",
        user_content="Current question: create an excel file of our plan",
    ), headers=free_user["headers"])
    assert r.json()["text"].startswith("If you were a Plus subscriber")


def test_at_or_above_min_tier_gets_no_upsell_line(
        client, pro_user, mock_provider, monkeypatch):
    """A pro user on a pro-gated feature rides the normal armed/teaser
    machinery — never the upsell."""
    from unittest.mock import AsyncMock
    import app.services.document_generation as dg
    from tests.conftest import chat_request

    _enable_generation_with_upsell(client)
    monkeypatch.setattr(dg, "classify_generation_intent", AsyncMock(
        return_value={"file_request": False, "format": None, "gist": ""}))
    r = client.post("/v1/chat", json=chat_request(
        prompt_mode="ProjectChat", call_type="query",
        user_content="Current question: summarize the meeting",
    ), headers=pro_user["headers"])
    assert "If you were a" not in r.json()["text"]


def test_below_tier_non_file_ask_gets_no_upsell_line(
        client, free_user, mock_provider):
    from tests.conftest import chat_request

    _enable_generation_with_upsell(client)
    r = client.post("/v1/chat", json=chat_request(
        prompt_mode="ProjectChat", call_type="query",
        user_content="Current question: summarize the last meeting",
    ), headers=free_user["headers"])
    assert r.json()["text"] == "Test response from mock provider."
    assert "If you were a" not in r.json()["text"]


def test_upsell_line_rides_the_stream_as_first_delta(
        client, free_user, monkeypatch):
    """Meeting chat streams tokens — the upsell line arrives as a
    synthetic first text delta in the same event shape the client
    already concatenates."""
    import json as _json
    from app.models.chat import ChatResponse
    from tests.conftest import chat_request

    _enable_generation_with_upsell(client)

    def _fake_stream(provider_router, body, db, settings):
        async def _gen():
            yield {"text": "Here is the plan in chat."}
            yield {"done": True, "response": ChatResponse(
                text="Here is the plan in chat.", input_tokens=10,
                output_tokens=5, model="claude-haiku-4-5-20251001",
                provider="anthropic",
                usage={"input_tokens": 10, "output_tokens": 5})}
        return _gen()

    monkeypatch.setattr(
        "app.services.anthropic_or_fallback.route_stream_with_fallback",
        _fake_stream)
    r = client.post("/v1/chat", json=chat_request(
        prompt_mode="PostMeetingChat", call_type="query", stream=True,
        user_content="Current question: create an excel file of our plan",
    ), headers=free_user["headers"])
    assert r.status_code == 200
    assert r.headers["content-type"].startswith("text/event-stream")
    texts = [_json.loads(line[len("data: "):])
             for line in r.text.splitlines()
             if line.startswith("data: ")]
    text_events = [e["text"] for e in texts if e.get("type") == "text"]
    assert text_events[0].startswith("If you were a Pro subscriber")
    assert text_events[1] == "Here is the plan in chat."


def test_upsell_text_is_locale_served(client, free_user, mock_provider):
    from tests.conftest import chat_request

    _enable_generation_with_upsell(client)
    es = client.app.state.remote_configs.setdefault("client-config.es", {})
    es.setdefault("documents", {})["generation"] = {
        "enabled": True, "min_tier": "pro",
        "upsell": {"enabled": True,
                   "text": "Con una suscripción {tier} podría generar el archivo."},
    }
    r = client.post("/v1/chat", json=chat_request(
        prompt_mode="ProjectChat", call_type="query",
        user_content="Current question: create an excel file of our plan",
    ), headers={**free_user["headers"], "Accept-Language": "es-MX"})
    assert r.json()["text"].startswith(
        "Con una suscripción Pro podría generar el archivo.")


def test_teaser_rides_the_stream_done_event(client, free_user, monkeypatch):
    """Meeting Chat device test 2026-07-14 (SS): the teaser was computed
    and its offer minted, but streaming surfaces never attached the
    envelope — the JSON tail was the only carrier. The teaser now rides
    the SSE done event (same vehicle as search_state), identical shape."""
    import json as _json
    from unittest.mock import AsyncMock
    import app.services.document_generation as dg
    from app.models.chat import ChatResponse
    from tests.conftest import chat_request

    _enable_confirmed_generation(client)
    monkeypatch.setattr(dg, "classify_generation_intent", AsyncMock(
        return_value={"file_request": False, "format": None, "gist": ""}))

    def _fake_stream(provider_router, body, db, settings):
        async def _gen():
            yield {"text": "A gantt would have phases and dates."}
            yield {"done": True, "response": ChatResponse(
                text="A gantt would have phases and dates.", input_tokens=10,
                output_tokens=5, model="claude-haiku-4-5-20251001",
                provider="anthropic",
                usage={"input_tokens": 10, "output_tokens": 5})}
        return _gen()

    monkeypatch.setattr(
        "app.services.anthropic_or_fallback.route_stream_with_fallback",
        _fake_stream)
    r = client.post("/v1/chat", json=chat_request(
        prompt_mode="PostMeetingChat", call_type="query", stream=True,
        user_content="Current question: what would a gantt chart of this "
                     "project look like?",
    ), headers=free_user["headers"])
    assert r.status_code == 200
    assert r.headers["content-type"].startswith("text/event-stream")
    events = [_json.loads(line[len("data: "):])
              for line in r.text.splitlines() if line.startswith("data: ")]
    done = next(e for e in events if e.get("type") == "done")
    fs = done.get("feature_state")
    assert fs and fs["cta"]["kind"] == "generation_teaser"
    assert fs["cta"]["details"]["offer_id"]  # typed-yes lane armed


def test_template_offer_gist_only_when_it_composes(
        client, free_user, mock_provider, monkeypatch):
    """Live 2026-07-14 (SS): classifier gist 'convert content to
    spreadsheet' jammed into 'Sounds like you want a project timeline
    convert content to spreadsheet.' Verb-phrase gists are dropped;
    qualifier gists still ride."""
    from unittest.mock import AsyncMock
    import app.services.document_generation as dg
    from tests.conftest import chat_request

    _enable_confirmed_generation(client)
    monkeypatch.setattr(dg, "classify_generation_intent", AsyncMock(
        return_value={"file_request": True, "format": "xlsx",
                      "gist": "convert content to spreadsheet"}))
    r = client.post("/v1/chat", json=chat_request(
        prompt_mode="ProjectChat", call_type="query",
        user_content="Current question: can you turn this gantt plan "
                     "into an Excel document?",
    ), headers=free_user["headers"])
    text = r.json()["feature_state"]["cta"]["text"]
    assert text.startswith("Sounds like you want a project timeline. ")
    assert "convert content" not in text
    assert "—" not in text  # served copy carries no em dashes

    monkeypatch.setattr(dg, "classify_generation_intent", AsyncMock(
        return_value={"file_request": True, "format": "xlsx",
                      "gist": "for the migration project"}))
    r2 = client.post("/v1/chat", json=chat_request(
        prompt_mode="ProjectChat", call_type="query",
        user_content="Current question: can you turn this gantt plan "
                     "into an Excel document?",
    ), headers=free_user["headers"])
    assert "a project timeline for the migration project." \
        in r2.json()["feature_state"]["cta"]["text"]


def test_template_format_veto_blocks_history_mismatch(
        client, free_user, mock_provider, monkeypatch):
    """Live 2026-07-14 21:58:42Z (Scott's device): 'make a word document
    that identifies the roles on today's call' drew the xlsx Gantt offer —
    'gantt' rode the carried conversation history, and the template's
    format never had to agree with the classifier's docx read. The FORMAT
    VETO closes it. History matching itself deliberately stays (anaphoric
    'make IT into excel' asks carry the keyword only in history — pinned
    by the surface-parity test)."""
    from unittest.mock import AsyncMock
    import app.services.document_generation as dg
    from tests.conftest import chat_request

    _enable_confirmed_generation(client)
    monkeypatch.setattr(dg, "classify_generation_intent", AsyncMock(
        return_value={"file_request": True, "format": "docx",
                      "gist": "for the roles on the call"}))
    r = client.post("/v1/chat", json=chat_request(
        prompt_mode="ProjectChat", call_type="query",
        user_content="Previous conversation: we discussed the gantt chart "
                     "and the project timeline milestones. Current question: "
                     "can you make a word document that identifies the "
                     "different roles and contributions of the people on "
                     "today's call?",
    ), headers=free_user["headers"])
    cta = r.json()["feature_state"]["cta"]
    assert "template_id" not in cta["details"]        # no gantt interception
    assert "Gantt" not in cta["text"]
    assert "Word document" in cta["text"] or "docx" in cta["text"]

    # format veto alone: gantt vocabulary IN the question, but the user
    # wants a Word file -> the xlsx template must not intercept
    monkeypatch.setattr(dg, "classify_generation_intent", AsyncMock(
        return_value={"file_request": True, "format": "docx", "gist": ""}))
    r2 = client.post("/v1/chat", json=chat_request(
        prompt_mode="ProjectChat", call_type="query",
        user_content="Current question: write a word doc summary of our "
                     "gantt timeline",
    ), headers=free_user["headers"])
    assert "template_id" not in r2.json()["feature_state"]["cta"]["details"]


def test_plain_offer_gist_composes_guard():
    """Second live jam 2026-07-14 22:25:27Z (Scott's device): plain offer
    read 'Sounds like you want a native Word document (.docx) seinfeld
    personality assignments for meeting attendees.' Noun-phrase gists now
    fall back to the no-gist offer copy; qualifier gists still ride."""
    from app.services.document_generation import (
        _CONFIRMATION_DEFAULTS,
        build_offer_envelope,
        gist_composes,
    )

    assert gist_composes("for onboarding new people") == "for onboarding new people"
    assert gist_composes("of the migration plan") == "of the migration plan"
    assert gist_composes("seinfeld personality assignments for attendees") == ""
    assert gist_composes("convert content to spreadsheet") == ""
    assert gist_composes("para el plan del proyecto") == "para el plan del proyecto"
    assert gist_composes(None) == "" and gist_composes("  ") == ""

    env = build_offer_envelope(
        _CONFIRMATION_DEFAULTS, "docx",
        gist="seinfeld personality assignments for meeting attendees")
    text = env["feature_state"]["cta"]["text"]
    assert text.startswith("That sounds like a file request.")  # no-gist copy
    assert "seinfeld" not in text
    assert env["feature_state"]["cta"]["details"]["gist"] == ""


def test_match_template_format_veto_unit():
    from app.services.doc_templates import match_template
    assert match_template("build a gantt chart") == "gantt_smartsheet"
    assert match_template("build a gantt chart", format="xlsx") == "gantt_smartsheet"
    assert match_template("build a gantt chart", format="docx") is None
    assert match_template("build a gantt chart", format=None) == "gantt_smartsheet"


def test_served_generation_copy_carries_no_em_dashes():
    """SS device test 2026-07-14: offer copy reached the screen with em
    dashes. Guard every served string source: code defaults, template
    offer nouns, and all three locale bundles."""
    from app.services.doc_templates import TEMPLATES
    from app.services.document_generation import (
        _CONFIRMATION_DEFAULTS,
        _UPSELL_DEFAULTS,
    )

    assert "—" not in json.dumps(_CONFIRMATION_DEFAULTS, ensure_ascii=False)
    assert "—" not in json.dumps(_UPSELL_DEFAULTS, ensure_ascii=False)
    for t in TEMPLATES.values():
        assert "—" not in t["offer_noun"]
    for f in ("client-config.json", "client-config.es.json",
              "client-config.ja.json"):
        gen = json.load(open(f"config/remote/{f}"))["documents"]["generation"]
        assert "—" not in json.dumps(gen, ensure_ascii=False), f


def test_bundled_upsell_live_all_locales():
    for f in ("client-config.json", "client-config.es.json",
              "client-config.ja.json"):
        up = json.load(open(f"config/remote/{f}"))["documents"]["generation"]["upsell"]
        assert up["enabled"] is True
        assert "{tier}" in up["text"]
        assert "—" not in up["text"]  # served copy carries no em dashes


def test_artifact_filename_is_distinctive():
    """Five identical Project_Gantt.xlsx rows in the References list made
    artifacts indistinguishable (Scott 2026-07-14): template artifacts now
    carry the extracted project slug and a MMDDYY stamp from the MEETING
    date (Scott's call), with the UTC build date only as fallback."""
    from app.services.doc_templates import TEMPLATES, artifact_filename
    t = TEMPLATES["gantt_smartsheet"]
    # meeting date from the plan wins (and makes the name deterministic)
    name = artifact_filename(t, {"project": "Kore Platform Rollout",
                                 "meeting_date": "2026-07-14"})
    assert name == "Kore_Platform_Rollout_Gantt_071426.xlsx"
    # no or broken meeting date -> UTC build-date fallback
    assert re.fullmatch(r"Gantt_\d{6}\.xlsx", artifact_filename(t, {}))
    assert re.fullmatch(
        r"Kore_Gantt_\d{6}\.xlsx",
        artifact_filename(t, {"project": "Kore", "meeting_date": "next tuesday"}))
    # hostile characters sanitize, long names cap
    name2 = artifact_filename(t, {"project": "a/b\\c: d*e?" + "x" * 100,
                                  "meeting_date": "2026-07-14"})
    assert re.fullmatch(r"a_b_c_d_e_x{30}_Gantt_071426\.xlsx", name2)
