"""Semantic docx rebuild — the Word-compat backstop for generated files.

Field finding (2026-07-11): the sandbox's first live .docx (authored with
docx.js) opened in every lenient reader — macOS importers, python-docx,
Google Docs — and hard-failed in Microsoft Word. Bisection localized the
poison to pervasive body markup that passes every schema check we can run,
i.e. only Word knows. The only intervention proven to fix it was a full
content rebuild (Scott's copy-paste-into-a-fresh-doc experiment).

This module is that experiment, programmatic: parse the artifact with
python-docx, reconstruct paragraphs / runs / tables / list styling onto a
fresh python-docx Document (whose template is derived from a real Word
file), and serialize with python-docx's writer. Deliberate fidelity trade:
exotic formatting is dropped; text, headings, bold/italic/underline,
alignment, bullet/number lists, and tables survive. A Word-openable file
with plain styling beats a pixel-faithful file Word refuses to open.

Fail-open: any rebuild error returns the ORIGINAL bytes — never lose the
artifact over the backstop.
"""

from __future__ import annotations

import logging
import re
from io import BytesIO

logger = logging.getLogger("ghostpour.docx_rebuild")

_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _numbering_formats(src) -> dict[str, str]:
    """numId -> level-0 numFmt ("bullet" / "decimal" / ...) from the source
    numbering part, so list paragraphs map to the right template style."""
    try:
        num_el = src.part.numbering_part.element
    except (AttributeError, KeyError, NotImplementedError):
        return {}
    abstract = {}
    for an in num_el.findall(f"{_W}abstractNum"):
        fmt = an.find(f"{_W}lvl/{_W}numFmt")
        abstract[an.get(f"{_W}abstractNumId")] = (
            fmt.get(f"{_W}val") if fmt is not None else "bullet")
    out = {}
    for num in num_el.findall(f"{_W}num"):
        ref = num.find(f"{_W}abstractNumId")
        if ref is not None:
            out[num.get(f"{_W}numId")] = abstract.get(ref.get(f"{_W}val"), "bullet")
    return out


def _copy_paragraph(p, out, num_formats: dict[str, str], out_styles: set[str]) -> None:
    style = None
    src_style = p.style.name if p.style is not None else None
    num_ref = p._p.find(f"{_W}pPr/{_W}numPr/{_W}numId")
    if num_ref is not None:
        # numbering-based list: map to the template's list styles by format
        fmt = num_formats.get(num_ref.get(f"{_W}val"), "bullet")
        style = "List Number" if fmt == "decimal" else "List Bullet"
    elif src_style and src_style != "Normal" and src_style in out_styles:
        # style-based markup (headings, style-based lists, quotes): keep any
        # style the Word-derived template also defines
        style = src_style
    new = out.add_paragraph(style=style)
    if p.alignment is not None:
        new.alignment = p.alignment
    for run in p.runs:
        r = new.add_run(run.text)
        r.bold, r.italic, r.underline = run.bold, run.italic, run.underline


def rebuild_docx(content: bytes) -> bytes:
    """Rebuild a .docx into python-docx-authored, Word-safe form.
    Returns the original bytes unchanged on any failure."""
    try:
        import docx
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        src = docx.Document(BytesIO(content))
        out = docx.Document()
        num_formats = _numbering_formats(src)
        out_styles = {s.name for s in out.styles}

        body = src.element.body
        blocks = 0
        for child in body:
            if child.tag == f"{_W}p":
                _copy_paragraph(Paragraph(child, src), out, num_formats, out_styles)
                blocks += 1
            elif child.tag == f"{_W}tbl":
                t = Table(child, src)
                cols = len(t.columns)
                new_t = out.add_table(rows=0, cols=cols)
                new_t.style = "Table Grid"
                for row in t.rows:
                    cells = new_t.add_row().cells
                    for i, cell in enumerate(row.cells[:cols]):
                        cells[i].text = "\n".join(
                            p.text for p in cell.paragraphs if p.text)
                blocks += 1
        # a fresh Document ships one empty leading paragraph — drop it
        first = out.paragraphs[0]
        if not first.text and len(out.paragraphs) > 1:
            first._p.getparent().remove(first._p)

        buf = BytesIO()
        out.save(buf)
        rebuilt = buf.getvalue()
        logger.info("docx_rebuild: %d blocks, %d -> %d bytes",
                    blocks, len(content), len(rebuilt))
        return rebuilt
    except Exception as e:
        logger.warning("docx_rebuild failed open — keeping original bytes: %s", e)
        return content
